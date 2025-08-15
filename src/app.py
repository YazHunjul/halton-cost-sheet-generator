"""
Main Streamlit application for the Halton Cost Sheet Generator.
"""
import streamlit as st
import os
import tempfile
import time
from datetime import datetime
from config.business_data import ESTIMATORS, SALES_CONTACTS, DELIVERY_LOCATIONS, COMPANY_ADDRESSES
from config.constants import VALID_CANOPY_MODELS
from utils.excel import read_excel_project_data, save_to_excel
from utils.word import generate_quotation_document
from utils.date_utils import format_date_for_display, get_current_date
from openpyxl import load_workbook
# from components.forms import general_project_form
# from components.project_forms import project_structure_form
# from config.constants import SessionKeys, PROJECT_TYPES
from utils.word import analyze_project_areas
from utils.state_manager import load_from_url, add_save_progress_button

def display_project_summary(project_data: dict):
    """Display a formatted summary of the project data."""
    st.header("Project Summary")
    
    # General Information
    st.subheader("General Information")
    col1, col2 = st.columns(2)
    
    with col1:
        st.write("**Project Name:**", project_data.get("project_name"))
        st.write("**Project Number:**", project_data.get("project_number"))
        st.write("**Date:**", format_date_for_display(project_data.get("date")))
        st.write("**Customer:**", project_data.get("customer"))
        st.write("**Company:**", project_data.get("company"))
    
    with col2:
        st.write("**Project Location:**", project_data.get("project_location") or project_data.get("location"))
        st.write("**Delivery Location:**", project_data.get("delivery_location"))
        st.write("**Address:**", project_data.get("address"))
        st.write("**Sales Contact:**", project_data.get("sales_contact"))
        st.write("**Estimator:**", project_data.get("estimator"))
    
    # Project Structure
    if "levels" in project_data:
        st.markdown("---")
        st.subheader("Project Structure")
        
        for level in project_data["levels"]:
            with st.expander(f"Level {level['level_number']}", expanded=True):
                for area in level["areas"]:
                    st.markdown(f"### 📍 Area: {area['name']}")
                    
                    # Area-level options
                    if "options" in area:
                        st.markdown("**Area Options:**")
                        opt_col1, opt_col2, opt_col3 = st.columns(3)
                        with opt_col1:
                            st.write("✓ UV-C System" if area["options"]["uvc"] else "✗ UV-C System")
                        with opt_col2:
                            st.write("✓ RecoAir" if area["options"]["recoair"] else "✗ RecoAir")
                        with opt_col3:
                            st.write("✓ Marvel" if area["options"]["marvel"] else "✗ Marvel")
                        st.markdown("---")
                    
                    if area["canopies"]:
                        for i, canopy in enumerate(area["canopies"], 1):
                            st.markdown(f"#### 🔹 Canopy {i}")
                            
                            # Basic Info
                            col1, col2 = st.columns(2)
                            with col1:
                                st.write("**Reference Number:**", canopy["reference_number"])
                                st.write("**Model:**", canopy["model"])
                                st.write("**Configuration:**", canopy["configuration"])
                            
                            # Wall Cladding
                            if canopy["wall_cladding"]["type"] != "None":
                                with col2:
                                    st.markdown("**Wall Cladding:**")
                                    st.write("- Type:", canopy["wall_cladding"]["type"])
                                    st.write("- Width:", f"{canopy['wall_cladding']['width']}mm")
                                    st.write("- Height:", f"{canopy['wall_cladding']['height']}mm")
                                    # Handle position as a list
                                    position = canopy["wall_cladding"]["position"]
                                    if isinstance(position, list):
                                        position_str = ", ".join(position) if position else "None"
                                    else:
                                        position_str = position if position else "None"
                                    st.write("- Position:", position_str)
                            
                            # Canopy Options (only fire suppression now)
                            st.markdown("**Canopy Options:**")
                            st.write("✓ Fire Suppression" if canopy["options"]["fire_suppression"] else "✗ Fire Suppression")
                            
                            st.markdown("---")
                    else:
                        st.write("No canopies in this area")
                    
                    st.markdown("---")

def word_generation_page():
    """Page for generating Word documents from uploaded Excel files."""
    st.header(" Generate Word Documents from Excel")
    st.markdown("Upload an existing Excel cost sheet to generate Word quotation documents.")
    
    uploaded_file = st.file_uploader(
        "Choose Excel file",
        type=['xlsx', 'xls'],
        help="Upload a cost sheet Excel file generated by this application"
    )
    
    if uploaded_file is not None:
        try:
            # Save uploaded file temporarily
            temp_path = f"temp_excel_{uploaded_file.name}"
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Read project data from Excel
            with st.spinner("Reading project data from Excel..."):
                project_data = read_excel_project_data(temp_path)
            
            # Display summary of extracted data
            st.success("Successfully extracted project data from Excel!")
            
            # Analyze project to show what type it is
            has_canopies, has_recoair, is_recoair_only, has_uv, has_marvel, has_vent_clg = analyze_project_areas(project_data)
            
            # Show project type analysis
            if is_recoair_only:
                st.info("**Project Type:** RecoAir-only project detected")
            elif has_canopies and has_recoair:
                st.info("**Project Type:** Mixed project (Canopies + RecoAir) detected")
            elif has_canopies:
                st.info("**Project Type:** Canopy-only project detected")
            else:
                st.warning("**Project Type:** No canopies or RecoAir systems detected")
            
            # Show download button first for quick access
            st.markdown("---")
            st.subheader("Quick Download")
            
            try:
                with st.spinner("Preparing download..."):
                    # Generate Word documents for download
                    download_word_path = generate_quotation_document(project_data, temp_path)
                
                # Provide download button
                if download_word_path.endswith('.zip'):
                    # Multiple documents in zip file
                    with open(download_word_path, "rb") as file:
                        zip_filename = os.path.basename(download_word_path)
                        st.download_button(
                            label="Download All Documents (ZIP)",
                            data=file.read(),
                            file_name=zip_filename,
                            mime="application/zip",
                            type="primary"
                        )
                    st.info("ZIP file contains both Main Quotation and RecoAir Quotation documents.")
                else:
                    # Single document
                    doc_filename = os.path.basename(download_word_path)
                    with open(download_word_path, "rb") as file:
                        # Determine appropriate label based on document type
                        if is_recoair_only:
                            label = "Download RecoAir Quotation"
                        else:
                            label = "Download Quotation"
                        
                        st.download_button(
                            label=label,
                            data=file.read(),
                            file_name=doc_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                    
                    # Show appropriate success message
                    if is_recoair_only:
                        st.info("RecoAir quotation document ready for download.")
                    else:
                        st.info("Quotation document ready for download.")
                        
            except Exception as e:
                st.error(f"Error preparing download: {str(e)}")
            
            # Automatically generate and show preview of existing document
            st.markdown("---")
            st.subheader("Document Preview")
            
            try:
                with st.spinner("Generating preview of current document..."):
                    # Generate Word documents to get the current state
                    word_path = generate_quotation_document(project_data, temp_path)
                
                # Show preview for documents
                if not word_path.endswith('.zip'):
                    # Single document preview
                    from utils.word_preview import check_preview_requirements, preview_word_document
                    capabilities = check_preview_requirements()
                    
                    col1, col2 = st.columns([3, 1])
                    with col2:
                        use_advanced = st.checkbox(
                            "Enhanced Preview", 
                            value=capabilities['advanced_preview'],
                            disabled=not capabilities['advanced_preview'],
                            help="Uses pypandoc for better formatting (if available)",
                            key="upload_preview_advanced"
                        )
                        
                        if not capabilities['advanced_preview']:
                            if "not installed" in str(capabilities.get('pandoc_version', '')):
                                st.info(" Install pypandoc for enhanced preview")
                            else:
                                st.warning(f" {capabilities.get('pandoc_version', 'Pandoc issue')}")
                        elif capabilities['pandoc_version']:
                            st.success(f"Yes Pandoc v{capabilities['pandoc_version']}")
                    
                    with col1:
                        st.write("**Preview Mode:**", "Enhanced" if use_advanced else "Basic")
                        if capabilities['table_preservation']:
                            st.write("Yes Table preservation enabled")
                    
                    # Generate and display preview
                    try:
                        with st.spinner("Rendering preview..."):
                            preview_html = preview_word_document(word_path, use_advanced)
                        
                        # Display preview
                        st.components.v1.html(preview_html, height=650, scrolling=True)
                        
                        # Preview stats
                        try:
                            from docx import Document
                            doc = Document(word_path)
                            table_count = len(doc.tables)
                            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
                            
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric(" Paragraphs", paragraph_count)
                            with col2:
                                st.metric(" Tables", table_count)
                            with col3:
                                file_size = os.path.getsize(word_path)
                                st.metric(" File Size", f"{file_size/1024:.1f} KB")
                                
                        except Exception as e:
                            st.write(f"Preview stats unavailable: {str(e)}")
                            
                    except Exception as e:
                        st.error(f"No Error generating preview: {str(e)}")
                        st.write("Preview failed, but you can still generate the document below.")
                else:
                    # Multiple documents - show previews for both
                    st.info(" Multiple documents detected - showing previews for both documents:")
                    
                    # Extract and preview individual documents from the ZIP
                    import zipfile
                    import tempfile
                    from utils.word_preview import check_preview_requirements, preview_word_document
                    
                    capabilities = check_preview_requirements()
                    
                    # Preview options (shared for both documents)
                    col1, col2 = st.columns([3, 1])
                    with col2:
                        use_advanced = st.checkbox(
                            "Enhanced Preview", 
                            value=capabilities['advanced_preview'],
                            disabled=not capabilities['advanced_preview'],
                            help="Uses pypandoc for better formatting (if available)",
                            key="upload_preview_advanced_multi"
                        )
                        
                        if not capabilities['advanced_preview']:
                            if "not installed" in str(capabilities.get('pandoc_version', '')):
                                st.info(" Install pypandoc for enhanced preview")
                            else:
                                st.warning(f" {capabilities.get('pandoc_version', 'Pandoc issue')}")
                        elif capabilities['pandoc_version']:
                            st.success(f"Yes Pandoc v{capabilities['pandoc_version']}")
                    
                    with col1:
                        st.write("**Preview Mode:**", "Enhanced" if use_advanced else "Basic")
                        if capabilities['table_preservation']:
                            st.write("Yes Table preservation enabled")
                    
                    try:
                        with zipfile.ZipFile(word_path, 'r') as zip_ref:
                            file_list = zip_ref.namelist()
                            
                            for i, filename in enumerate(file_list):
                                if filename.endswith('.docx'):
                                    st.markdown(f"###  Document {i+1}: {filename}")
                                    
                                    # Extract to temporary file
                                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                                        tmp_file.write(zip_ref.read(filename))
                                        tmp_path = tmp_file.name
                                    
                                    try:
                                        # Generate and display preview
                                        with st.spinner(f"Rendering preview for {filename}..."):
                                            preview_html = preview_word_document(tmp_path, use_advanced)
                                        
                                        # Display preview
                                        st.components.v1.html(preview_html, height=500, scrolling=True)
                                        
                                        # Preview stats
                                        try:
                                            from docx import Document
                                            doc = Document(tmp_path)
                                            table_count = len(doc.tables)
                                            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
                                            
                                            col1, col2, col3 = st.columns(3)
                                            with col1:
                                                st.metric(" Paragraphs", paragraph_count)
                                            with col2:
                                                st.metric(" Tables", table_count)
                                            with col3:
                                                file_size = os.path.getsize(tmp_path)
                                                st.metric(" File Size", f"{file_size/1024:.1f} KB")
                                                
                                        except Exception as e:
                                            st.write(f"Preview stats unavailable: {str(e)}")
                                    
                                    except Exception as e:
                                        st.error(f"No Error generating preview for {filename}: {str(e)}")
                                    
                                    finally:
                                        # Clean up temp file
                                        if os.path.exists(tmp_path):
                                            os.unlink(tmp_path)
                                    
                                    if i < len(file_list) - 1:  # Add separator between documents
                                        st.markdown("---")
                    
                    except Exception as e:
                        st.error(f"No Error extracting documents from ZIP: {str(e)}")
                        st.write("Preview failed, but you can still generate the documents below.")
                    
            except Exception as e:
                st.error(f"No Error generating preview: {str(e)}")
                st.write("Preview failed, but you can still generate the document below.")
            
            st.markdown("---")
            
            with st.expander(" Extracted Project Data", expanded=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.write("**Project Name:**", project_data.get("project_name"))
                    st.write("**Project Number:**", project_data.get("project_number"))
                    st.write("**Customer:**", project_data.get("customer"))
                    st.write("**Date:**", format_date_for_display(project_data.get("date")))
                
                with col2:
                    st.write("**Project Location:**", project_data.get("project_location") or project_data.get("location"))
                    st.write("**Delivery Location:**", project_data.get("delivery_location"))
                    st.write("**Estimator:**", project_data.get("estimator"))
                    st.write("**Estimator Initials (from Excel):**", project_data.get("estimator_initials"))
                    
                    # Show combined initials calculation, reference variable, customer first name, and quote title
                    from utils.word import get_sales_contact_info, get_combined_initials, generate_reference_variable, get_customer_first_name, generate_quote_title
                    estimator_name = project_data.get("estimator", "")
                    sales_contact = get_sales_contact_info(estimator_name, project_data)
                    combined_initials = get_combined_initials(sales_contact['name'], estimator_name)
                    reference_variable = generate_reference_variable(
                        project_data.get('project_number', ''), 
                        sales_contact['name'], 
                        estimator_name
                    )
                    customer_first_name = get_customer_first_name(project_data.get('customer', ''))
                    quote_title = generate_quote_title(project_data.get('revision', ''))
                    st.write("**Combined Initials (Sales/Estimator):**", combined_initials)
                    st.write("**Reference Variable:**", reference_variable)
                    st.write("**Customer First Name:**", customer_first_name)
                    st.write("**Quote Title:**", quote_title)
                    st.write("**Revision:**", project_data.get('revision', '') or 'Initial Version')
                    st.write("**Sales Contact:**", sales_contact['name'])
                    
                    st.write("**Levels Found:**", len(project_data.get("levels", [])))
                
                # Show detailed analysis
                st.markdown("---")
                st.markdown("** Project Analysis:**")
                analysis_col1, analysis_col2, analysis_col3 = st.columns(3)
                with analysis_col1:
                    st.write("**Has Canopies:**", "Yes" if has_canopies else "No")
                with analysis_col2:
                    st.write("**Has RecoAir:**", "Yes" if has_recoair else "No")
                with analysis_col3:
                    st.write("**RecoAir Only:**", "Yes" if is_recoair_only else "No")
                
                # Show areas and their options
                if project_data.get("levels"):
                    st.markdown("** Areas Found:**")
                    for level in project_data.get("levels", []):
                        for area in level.get("areas", []):
                            area_name = f"{level.get('level_name', '')} - {area.get('name', '')}"
                            canopy_count = len(area.get('canopies', []))
                            options = area.get('options', {})
                            
                            st.write(f"• **{area_name}**: {canopy_count} canopies")
                            if options.get('uvc'):
                                st.write("  - Yes UV-C System")
                            if options.get('recoair'):
                                st.write("  - Yes RecoAir System")
            
            # Show what documents will be generated
            st.markdown("---")
            st.markdown("** Documents to Generate:**")
            if is_recoair_only:
                st.info(" **RecoAir Quotation** will be generated (single document)")
                st.write(" Your Excel file now has dynamic pricing - totals update automatically!")
            elif has_canopies and has_recoair:
                st.info(" **ZIP Package** will be generated containing:")
                st.write("• Main Quotation (for canopies)")
                st.write("• RecoAir Quotation (for RecoAir systems)")
                st.write(" Your Excel file now has dynamic pricing - totals update automatically!")
            elif has_canopies:
                st.info(" **Main Quotation** will be generated (single document)")
                st.write(" Your Excel file now has dynamic pricing - totals update automatically!")
            else:
                st.warning(" No documents can be generated - no systems detected")
            
            # Generate Word document
            if st.button(" Generate Word Quotation", type="primary"):
                try:
                    with st.spinner("Generating Word quotation document(s)..."):
                        # Generate Word documents only (Excel has dynamic pricing now)
                        word_path = generate_quotation_document(project_data, temp_path)
                    
                    st.success("Yes Word quotation document(s) generated successfully!")
                    
                    # Determine file type and provide appropriate download button
                    if word_path.endswith('.zip'):
                        # Multiple documents in zip file
                        with open(word_path, "rb") as file:
                            # Extract filename from the generated path
                            zip_filename = os.path.basename(word_path)
                            st.download_button(
                                label=" Download Quotation Documents (ZIP)",
                                data=file.read(),
                                file_name=zip_filename,
                                mime="application/zip"
                            )
                        st.info(" Multiple quotation documents generated and packaged in ZIP file.")
                    else:
                        # Single document - automatically show preview with download option
                        doc_filename = os.path.basename(word_path)
                        
                        # Determine appropriate success message based on document type
                        if is_recoair_only:
                            st.info(" RecoAir quotation document generated successfully.")
                        else:
                            st.info(" Quotation document generated successfully.")
                        
                        # Show download button first
                        with open(word_path, "rb") as file:
                            # Determine appropriate label based on document type
                            if is_recoair_only:
                                label = " Download RecoAir Quotation"
                            else:
                                label = " Download Quotation"
                            
                            st.download_button(
                                label=label,
                                data=file.read(),
                                file_name=doc_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        # Automatically show preview below
                        st.markdown("---")
                        from utils.word_preview import preview_with_download
                        
                        # Show preview with a more compact interface
                        st.subheader(" Document Preview")
                        
                        # Check capabilities and show preview options
                        from utils.word_preview import check_preview_requirements
                        capabilities = check_preview_requirements()
                        
                        col1, col2 = st.columns([3, 1])
                        with col2:
                            use_advanced = st.checkbox(
                                "Enhanced Preview", 
                                value=capabilities['advanced_preview'],
                                disabled=not capabilities['advanced_preview'],
                                help="Uses pypandoc for better formatting (if available)"
                            )
                            
                            if not capabilities['advanced_preview']:
                                st.info(" Install pypandoc for enhanced preview")
                            elif capabilities['pandoc_version']:
                                st.caption(f"Pandoc version: {capabilities['pandoc_version']}")
                        
                        with col1:
                            st.write("**Preview Mode:**", "Enhanced" if use_advanced else "Basic")
                            if capabilities['table_preservation']:
                                st.write("Yes Table preservation enabled")
                        
                        # Generate and display preview
                        try:
                            from utils.word_preview import preview_word_document
                            with st.spinner("Generating preview..."):
                                preview_html = preview_word_document(word_path, use_advanced)
                            
                            # Display preview
                            st.components.v1.html(preview_html, height=650, scrolling=True)
                            
                            # Preview stats
                            try:
                                from docx import Document
                                doc = Document(word_path)
                                table_count = len(doc.tables)
                                paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
                                
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.metric(" Paragraphs", paragraph_count)
                                with col2:
                                    st.metric(" Tables", table_count)
                                with col3:
                                    file_size = os.path.getsize(word_path)
                                    st.metric(" File Size", f"{file_size/1024:.1f} KB")
                                    
                            except Exception as e:
                                st.write(f"Preview stats unavailable: {str(e)}")
                                
                        except Exception as e:
                            st.error(f"No Error generating preview: {str(e)}")
                            st.write("The document was generated successfully but preview failed. You can still download it above.")
                    
                    # Add note about dynamic pricing
                    st.success("✨ **Your Excel file now has dynamic pricing!** The JOB TOTAL sheet will automatically update when you edit any individual sheet prices.")
                
                except Exception as e:
                    st.error(f"No Error generating Word document: {str(e)}")
            
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
        except Exception as e:
            error_message = str(e)
            
            # Check if this is a validation error with detailed information
            if "Data validation errors found:" in error_message:
                st.error("❌ **Excel File Validation Errors**")
                st.markdown("The following data validation errors were found in your Excel file:")
                
                # Split the error message to extract the validation details
                parts = error_message.split("Data validation errors found:")
                if len(parts) > 1:
                    validation_details = parts[1].strip()
                    # Display each validation error in an expandable section
                    with st.expander(" **Detailed Error Information**", expanded=True):
                        st.markdown(validation_details)
                
                st.markdown("---")
                st.markdown("###  **How to Fix:**")
                st.markdown("1. Open your Excel file")
                st.markdown("2. Navigate to the specific cells mentioned above")
                st.markdown("3. Ensure all numeric fields contain valid numbers (not letters or text)")
                st.markdown("4. Save the file and try uploading again")
                
                st.info(" **Tip:** The most common issue is entering letters in numeric fields like 'Testing and Commissioning' prices.")
                
            else:
                st.error(f"❌ Error reading Excel file: {error_message}")
                # Show detailed traceback for debugging
                with st.expander("🔍 Technical Details", expanded=False):
                    import traceback
                    st.code(traceback.format_exc())
            
            if os.path.exists(temp_path):
                os.remove(temp_path)

def revision_page():
    """Page for creating new revisions from existing Excel files with full editing capabilities."""
    st.header(" Create & Edit Revision")
    st.markdown("Upload an existing Excel cost sheet to create a new revision. You can edit all aspects of the project before generating the revision.")
    
    uploaded_file = st.file_uploader(
        "Choose Excel file to revise",
        type=['xlsx', 'xls'],
        help="Upload an existing cost sheet Excel file to create a new revision"
    )
    
    if uploaded_file is not None:
        try:
            # Save uploaded file temporarily
            temp_path = f"temp_revision_{uploaded_file.name}"
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Read project data from Excel
            with st.spinner("Reading project data from Excel..."):
                project_data = read_excel_project_data(temp_path)
            
            # Display summary of extracted data
            st.success("✅ Successfully extracted project data from Excel!")
            
            # Show current revision info
            current_revision = project_data.get('revision', '')
            revision_display = current_revision if current_revision else 'Initial Version'
            
            # Reset session state for new file upload (use filename as key to detect changes)
            current_file_key = f"{uploaded_file.name}_{uploaded_file.size}"
            if 'revision_file_key' not in st.session_state or st.session_state.revision_file_key != current_file_key:
                # Clear old data
                if 'revision_project_data' in st.session_state:
                    del st.session_state.revision_project_data
                if 'revision_levels' in st.session_state:
                    del st.session_state.revision_levels
                st.session_state.revision_file_key = current_file_key
            
            # Store project data in session state for editing
            if 'revision_project_data' not in st.session_state:
                st.session_state.revision_project_data = project_data.copy()
                # Convert levels data to match the format used in the main editor
                if 'levels' in project_data:
                    # Convert from Excel format to editor format
                    st.session_state.revision_levels = []
                    try:
                        for level in project_data['levels']:
                            converted_level = {
                                'name': level.get('level_name', f"Level {level.get('level_number', '')}"),
                                'areas': []
                            }
                            # Safely process areas
                            if 'areas' in level and level['areas']:
                                for area in level['areas']:
                                    # Ensure each area has required fields
                                    if isinstance(area, dict):
                                        converted_level['areas'].append(area)
                            st.session_state.revision_levels.append(converted_level)
                    except Exception as e:
                        st.error(f"Error processing levels data: {str(e)}")
                        st.write("Level structure that caused error:")
                        st.json(project_data['levels'])
                        st.session_state.revision_levels = []
                else:
                    st.session_state.revision_levels = []
            
            # Tabs for different editing sections
            tab1, tab2, tab3, tab4 = st.tabs([
                " Project Info", 
                " Levels & Areas", 
                " Canopy Configuration",
                " Generate Revision"
            ])
            
            with tab1:
                st.subheader("Edit Project Information")
                col1, col2 = st.columns(2)
                
                with col1:
                    st.session_state.revision_project_data['project_name'] = st.text_input(
                        "Project Name",
                        value=st.session_state.revision_project_data.get('project_name', ''),
                        key="rev_project_name"
                    )
                    st.session_state.revision_project_data['project_number'] = st.text_input(
                        "Project Number",
                        value=st.session_state.revision_project_data.get('project_number', ''),
                        key="rev_project_number"
                    )
                    st.session_state.revision_project_data['customer'] = st.text_input(
                        "Customer",
                        value=st.session_state.revision_project_data.get('customer', ''),
                        key="rev_customer"
                    )
                    st.session_state.revision_project_data['company'] = st.text_input(
                        "Company",
                        value=st.session_state.revision_project_data.get('company', ''),
                        key="rev_company"
                    )
                
                with col2:
                    st.session_state.revision_project_data['project_location'] = st.text_input(
                        "Project Location",
                        value=st.session_state.revision_project_data.get('project_location') or st.session_state.revision_project_data.get('location', ''),
                        key="rev_project_location"
                    )
                    
                    # Delivery location dropdown
                    from config.business_data import DELIVERY_LOCATIONS
                    current_delivery = st.session_state.revision_project_data.get('delivery_location', '')
                    if current_delivery and current_delivery not in DELIVERY_LOCATIONS:
                        delivery_options = [current_delivery] + DELIVERY_LOCATIONS
                    else:
                        delivery_options = DELIVERY_LOCATIONS
                    
                    st.session_state.revision_project_data['delivery_location'] = st.selectbox(
                        "Delivery Location",
                        options=delivery_options,
                        index=delivery_options.index(current_delivery) if current_delivery in delivery_options else 0,
                        key="rev_delivery_location"
                    )
                    
                    st.session_state.revision_project_data['estimator'] = st.text_input(
                        "Estimator",
                        value=st.session_state.revision_project_data.get('estimator', ''),
                        key="rev_estimator"
                    )
                    
                    # Show current revision (read-only)
                    st.text_input(
                        "Current Revision",
                        value=revision_display,
                        disabled=True,
                        key="rev_current_revision"
                    )
            
            with tab2:
                st.subheader("Edit Levels & Areas")
                
                # Function to edit levels and areas for revision
                if 'revision_levels' not in st.session_state:
                    st.session_state.revision_levels = []
                
                # Add new level button
                if st.button(" Add New Level", key="rev_add_level"):
                    new_level = {
                        'name': f"Level {len(st.session_state.revision_levels) + 1}",
                        'areas': []
                    }
                    st.session_state.revision_levels.append(new_level)
                    st.rerun()
                
                # Display and edit existing levels
                for level_idx, level in enumerate(st.session_state.revision_levels):
                    # Ensure level has a name
                    if 'name' not in level:
                        level['name'] = level.get('level_name', f"Level {level_idx + 1}")
                    
                    with st.expander(f" {level['name']}", expanded=True):
                        col1, col2 = st.columns([3, 1])
                        
                        with col1:
                            # Edit level name
                            level['name'] = st.text_input(
                                "Level Name",
                                value=level['name'],
                                key=f"rev_level_name_{level_idx}"
                            )
                        
                        with col2:
                            # Remove level button
                            if st.button("No Remove Level", key=f"rev_remove_level_{level_idx}"):
                                st.session_state.revision_levels.pop(level_idx)
                                st.rerun()
                        
                        # Areas for this level
                        st.write("**Areas:**")
                        
                        # Add area button
                        if st.button(f" Add Area to {level['name']}", key=f"rev_add_area_{level_idx}"):
                            new_area = {
                                'name': f"Area {len(level['areas']) + 1}",
                                'canopies': []
                            }
                            level['areas'].append(new_area)
                            st.rerun()
                        
                        # Display areas
                        for area_idx, area in enumerate(level['areas']):
                            # Ensure area has required fields
                            if 'name' not in area:
                                area['name'] = f"Area {area_idx + 1}"
                            if 'options' not in area:
                                area['options'] = {}
                            
                            with st.container():
                                col1, col2 = st.columns([3, 1])
                                
                                with col1:
                                    area['name'] = st.text_input(
                                        f"Area Name",
                                        value=area.get('name', f"Area {area_idx + 1}"),
                                        key=f"rev_area_name_{level_idx}_{area_idx}"
                                    )
                                
                                with col2:
                                    if st.button("❌ Remove", key=f"rev_remove_area_{level_idx}_{area_idx}"):
                                        level['areas'].pop(area_idx)
                                        st.rerun()
                                
                                # Area options
                                st.write("**Area Options:**")
                                opt_col1, opt_col2, opt_col3, opt_col4 = st.columns(4)
                                
                                with opt_col1:
                                    area['options']['uvc'] = st.checkbox(
                                        "UV-C",
                                        value=area.get('options', {}).get('uvc', False),
                                        key=f"rev_area_uvc_{level_idx}_{area_idx}"
                                    )
                                
                                with opt_col2:
                                    area['options']['recoair'] = st.checkbox(
                                        "RecoAir",
                                        value=area.get('options', {}).get('recoair', False),
                                        key=f"rev_area_recoair_{level_idx}_{area_idx}"
                                    )
                                
                                with opt_col3:
                                    area['options']['marvel'] = st.checkbox(
                                        "Marvel",
                                        value=area.get('options', {}).get('marvel', False),
                                        key=f"rev_area_marvel_{level_idx}_{area_idx}"
                                    )
                                
                                with opt_col4:
                                    area['options']['vent_clg'] = st.checkbox(
                                        "Vent CLG",
                                        value=area.get('options', {}).get('vent_clg', False),
                                        key=f"rev_area_vent_{level_idx}_{area_idx}"
                                    )
                                
                                st.markdown("---")  # Separator between areas
            
            with tab3:
                st.subheader("Edit Canopy Configuration")
                
                if not st.session_state.revision_levels:
                    st.info("Please add levels and areas in the Levels & Areas tab first.")
                else:
                    # Display canopy configuration for each area
                    for level_idx, level in enumerate(st.session_state.revision_levels):
                        if level['areas']:
                            st.write(f"### {level['name']}")
                            
                            for area_idx, area in enumerate(level['areas']):
                                # Ensure area has a name
                                area_name = area.get('name', f"Area {area_idx + 1}")
                                with st.expander(f" {area_name}", expanded=True):
                                    # Add canopy button
                                    if st.button(f"➕ Add Canopy", key=f"rev_add_canopy_{level_idx}_{area_idx}"):
                                        if 'canopies' not in area:
                                            area['canopies'] = []
                                        new_canopy = {
                                            "reference_number": f"C{len(area['canopies']) + 1:03d}",
                                            "configuration": "",
                                            "model": "",
                                            "length": 1000,
                                            "width": 1000,
                                            "height": 555,
                                            "sections": 1,
                                            "options": {
                                                "fire_suppression": False,
                                                "sdu": False
                                            },
                                            "wall_cladding": {
                                                "type": "None",
                                                "width": 0,
                                                "height": 2100,
                                                "position": []
                                            },
                                            "sdu_item_number": ""
                                        }
                                        area['canopies'].append(new_canopy)
                                        st.rerun()
                                    
                                    # Display existing canopies
                                    if 'canopies' in area and area['canopies']:
                                        for canopy_idx, canopy in enumerate(area['canopies']):
                                            canopy_key = f"rev_canopy_{level_idx}_{area_idx}_{canopy_idx}"
                                            
                                            with st.container():
                                                # Header with remove button
                                                header_col1, header_col2 = st.columns([5, 1])
                                                with header_col1:
                                                    st.write(f"**Canopy {canopy_idx + 1} - {canopy.get('reference_number', f'C{canopy_idx + 1:03d}')}**")
                                                with header_col2:
                                                    if st.button("❌", key=f"{canopy_key}_remove"):
                                                        area['canopies'].pop(canopy_idx)
                                                        st.rerun()
                                                
                                                # Basic info
                                                col1, col2, col3, col4 = st.columns(4)
                                                
                                                with col1:
                                                    canopy['reference_number'] = st.text_input(
                                                        "Reference No.",
                                                        value=canopy.get('reference_number', f'C{canopy_idx + 1:03d}'),
                                                        key=f"{canopy_key}_ref"
                                                    )
                                                    
                                                with col2:
                                                    from config.business_data import VALID_CANOPY_MODELS
                                                    model_options = [""] + VALID_CANOPY_MODELS
                                                    canopy['model'] = st.selectbox(
                                                        "Model",
                                                        options=model_options,
                                                        index=0 if canopy.get('model', '') == '' else (model_options.index(canopy.get('model', '')) if canopy.get('model', '') in model_options else 0),
                                                        key=f"{canopy_key}_model"
                                                    )
                                                
                                                with col3:
                                                    # Configuration can be edited
                                                    configuration_options = ["WALL", "ISLAND"]
                                                    current_config = canopy.get('configuration', '')
                                                    if current_config and current_config not in configuration_options:
                                                        configuration_options.insert(0, current_config)
                                                    canopy['configuration'] = st.selectbox(
                                                        "Configuration",
                                                        options=configuration_options,
                                                        index=configuration_options.index(current_config) if current_config in configuration_options else 0,
                                                        key=f"{canopy_key}_config"
                                                    )
                                                
                                                with col4:
                                                    canopy['sections'] = st.number_input(
                                                        "Sections",
                                                        value=int(canopy.get('sections', 1)),
                                                        min_value=1,
                                                        max_value=10,
                                                        key=f"{canopy_key}_sections"
                                                    )
                                                
                                                # Dimensions
                                                col1, col2, col3 = st.columns(3)
                                                
                                                with col1:
                                                    canopy['length'] = st.number_input(
                                                        "Length (mm)",
                                                        value=int(canopy.get('length', 1000)),
                                                        min_value=0,
                                                        step=100,
                                                        key=f"{canopy_key}_length"
                                                    )
                                                
                                                with col2:
                                                    canopy['width'] = st.number_input(
                                                        "Width (mm)",
                                                        value=int(canopy.get('width', 1000)),
                                                        min_value=0,
                                                        step=100,
                                                        key=f"{canopy_key}_width"
                                                    )
                                                
                                                with col3:
                                                    canopy['height'] = st.number_input(
                                                        "Height (mm)",
                                                        value=int(canopy.get('height', 555)),
                                                        min_value=0,
                                                        step=50,
                                                        key=f"{canopy_key}_height"
                                                    )
                                                
                                                # Options
                                                st.write("**Canopy Options:**")
                                                opt_col1, opt_col2, opt_col3, opt_col4 = st.columns(4)
                                                
                                                with opt_col1:
                                                    if 'options' not in canopy:
                                                        canopy['options'] = {}
                                                    canopy['options']['fire_suppression'] = st.checkbox(
                                                        "Fire Suppression",
                                                        value=canopy.get('options', {}).get('fire_suppression', False),
                                                        key=f"{canopy_key}_fire"
                                                    )
                                                
                                                with opt_col2:
                                                    canopy['options']['sdu'] = st.checkbox(
                                                        "SDU",
                                                        value=canopy.get('options', {}).get('sdu', False),
                                                        key=f"{canopy_key}_sdu"
                                                    )
                                                    
                                                with opt_col3:
                                                    # If SDU is selected, show item number
                                                    if canopy['options']['sdu']:
                                                        canopy['sdu_item_number'] = st.text_input(
                                                            "SDU Item No.",
                                                            value=canopy.get('sdu_item_number', ''),
                                                            key=f"{canopy_key}_sdu_item"
                                                        )
                                                
                                                # Wall Cladding
                                                st.write("**Wall Cladding:**")
                                                if 'wall_cladding' not in canopy:
                                                    canopy['wall_cladding'] = {"type": "None", "width": 0, "height": 0, "position": []}
                                                
                                                # Check if wall cladding is enabled (not 'None' and has some data)
                                                wall_cladding_data = canopy.get('wall_cladding', {})
                                                has_wall_cladding = (
                                                    wall_cladding_data.get('type', 'None') not in ['None', None] or
                                                    wall_cladding_data.get('width', 0) or
                                                    wall_cladding_data.get('height', 0) or
                                                    wall_cladding_data.get('position', [])
                                                )
                                                
                                                wall_clad_enabled = st.checkbox(
                                                    "With Wall Cladding",
                                                    value=has_wall_cladding,
                                                    key=f"{canopy_key}_wall_clad"
                                                )
                                                
                                                if wall_clad_enabled:
                                                    clad_col1, clad_col2, clad_col3 = st.columns(3)
                                                    
                                                    # Set type to Stainless Steel by default (no UI input)
                                                    canopy['wall_cladding']['type'] = 'Stainless Steel'
                                                    
                                                    with clad_col1:
                                                        # Handle None values for width
                                                        width_value = canopy['wall_cladding'].get('width', 0)
                                                        if width_value is None:
                                                            width_value = 0
                                                        canopy['wall_cladding']['width'] = st.number_input(
                                                            "Width (mm)",
                                                            value=int(width_value),
                                                            min_value=0,
                                                            step=100,
                                                            key=f"{canopy_key}_clad_width"
                                                        )
                                                    
                                                    with clad_col2:
                                                        # Handle None values for height, default to 2100
                                                        height_value = canopy['wall_cladding'].get('height', 2100)
                                                        if height_value is None or height_value == 0:
                                                            height_value = 2100
                                                        canopy['wall_cladding']['height'] = st.number_input(
                                                            "Height (mm)",
                                                            value=int(height_value),
                                                            min_value=0,
                                                            step=100,
                                                            key=f"{canopy_key}_clad_height"
                                                        )
                                                    
                                                    with clad_col3:
                                                        cladding_positions = ["rear", "left hand", "right hand"]
                                                        position_value = canopy['wall_cladding'].get('position', [])
                                                        if isinstance(position_value, str):
                                                            position_value = [position_value]
                                                        elif not isinstance(position_value, list):
                                                            position_value = []
                                                        
                                                        selected_positions = st.multiselect(
                                                            "Position",
                                                            options=cladding_positions,
                                                            default=position_value,
                                                            key=f"{canopy_key}_clad_pos"
                                                        )
                                                        canopy['wall_cladding']['position'] = selected_positions
                                                else:
                                                    canopy['wall_cladding'] = {"type": "None", "width": 0, "height": 2100, "position": []}
                                                
                                                st.markdown("---")  # Separator between canopies
            
            with tab4:
                st.subheader(" Generate Revision")
                
                # Auto-increment revision
                if current_revision == '':
                    next_revision = 'A'
                else:
                    next_revision = chr(ord(current_revision) + 1) if current_revision and len(current_revision) == 1 and current_revision < 'Z' else 'B'
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write(f"**Current Revision:** {revision_display}")
                    st.write(f"**Suggested Next Revision:** {next_revision}")
                    
                    revision_choice = st.radio(
                        "Choose revision method:",
                        ["Auto-increment (recommended)", "Custom revision letter"],
                        help="Auto-increment will automatically suggest the next revision letter"
                    )
                
                with col2:
                    if revision_choice == "Custom revision letter":
                        new_revision = st.text_input(
                            "Enter new revision letter:",
                            value=next_revision,
                            max_chars=3,
                            help="Enter a revision letter (e.g., B, C, D, etc.)"
                        ).upper()
                    else:
                        new_revision = next_revision
                        st.write(f"**New Revision will be:** {new_revision}")
                    
                    # Update date option
                    update_date = st.checkbox("Update date to today", value=True)
                    
                    if update_date:
                        new_date = get_current_date()
                        st.write(f"**New Date:** {new_date}")
                    else:
                        new_date = st.session_state.revision_project_data.get("date", "")
                        st.write(f"**Date will remain:** {format_date_for_display(new_date)}")
                
                # Contract sheets option
                st.markdown("---")
                st.subheader("📋 Contract Options")
                
                # Initialize contract option in session state if not present
                if 'revision_contract_option' not in st.session_state:
                    st.session_state.revision_contract_option = st.session_state.revision_project_data.get('contract_option', False)
                
                include_contract_sheets = st.checkbox(
                    "Include Contract Sheets",
                    value=st.session_state.revision_contract_option,
                    key="rev_contract_checkbox",
                    help="Include CONTRACT, EXTRACT DUCT, SUPPLY DUCT, and SPIRAL DUCT sheets in the Excel file"
                )
                st.session_state.revision_contract_option = include_contract_sheets
                
                if include_contract_sheets:
                    st.info("✅ Contract sheets will be included in the revision")
                else:
                    st.info("ℹ️ Contract sheets will not be included in the revision")
                
                st.markdown("---")
                
                # Generate button
                if st.button(f" Generate Revision {new_revision}", type="primary", use_container_width=True):
                    try:
                        with st.spinner(f"Generating revision {new_revision} with your edits..."):
                            # Convert levels back to Excel format
                            converted_levels = []
                            for idx, level in enumerate(st.session_state.revision_levels):
                                converted_level = {
                                    'level_number': idx + 1,
                                    'level_name': level['name'],
                                    'areas': level.get('areas', [])
                                }
                                converted_levels.append(converted_level)
                            
                            # Update the project data with edited levels
                            st.session_state.revision_project_data['levels'] = converted_levels
                            st.session_state.revision_project_data['revision'] = new_revision
                            if update_date:
                                st.session_state.revision_project_data['date'] = new_date
                            
                            # Update contract option
                            st.session_state.revision_project_data['contract_option'] = include_contract_sheets
                            
                            # Create revision by properly regenerating the Excel file with all changes
                            # This ensures all canopy additions, modifications, and other changes are saved
                            from utils.excel import save_to_excel
                            
                            # Determine the template to use based on original file or default to latest
                            template_used = st.session_state.revision_project_data.get('template_used', 'R19.2')
                            if template_used == 'R19.1':
                                template_path = 'templates/excel/Cost Sheet R19.1 May 2025.xlsx'
                            elif template_used == 'R18.1':
                                template_path = 'templates/excel/Halton Cost Sheet Jan 2025.xlsx'
                            else:
                                template_path = 'templates/excel/Cost Sheet R19.2 Jun 2025.xlsx'  # Default to latest
                            
                            # Generate the Excel file with all the edited data
                            output_path = save_to_excel(
                                st.session_state.revision_project_data,
                                template_path=template_path
                            )
                            
                            # Read the file for download
                            with open(output_path, "rb") as file:
                                excel_data = file.read()
                        
                        st.success(f"Yes Revision {new_revision} created successfully with all your edits!")
                        
                        # Create download filename
                        project_number = st.session_state.revision_project_data.get('project_number', 'unknown')
                        date_str = new_date.replace('/', '') if update_date else st.session_state.revision_project_data.get('date', '').replace('/', '')
                        download_filename = f"{project_number} Cost Sheet {date_str}.xlsx"
                        
                        st.download_button(
                            label=f" Download Revision {new_revision}",
                            data=excel_data,
                            file_name=download_filename,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                        
                        # Show summary of changes
                        st.info(" **Summary of Changes:**")
                        st.write(f"• Revision updated: {current_revision} → {new_revision}")
                        if update_date:
                            st.write(f"• Date updated to: {new_date}")
                        st.write(f"• Total levels: {len(st.session_state.revision_levels)}")
                        total_areas = sum(len(level['areas']) for level in st.session_state.revision_levels)
                        st.write(f"• Total areas: {total_areas}")
                        st.write(f"• Contract sheets: {'Included' if include_contract_sheets else 'Not included'}")
                        st.write("• Yes All edits have been applied")
                        st.write("• Yes All existing data preserved (lights, formulas, etc.)")
                        st.write("• Yes Only edited fields were updated")
                        
                    except Exception as e:
                        st.error(f"❌ Error creating revision: {str(e)}")
            
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
        except Exception as e:
            error_message = str(e)
            
            # Check if this is a validation error with detailed information
            if "Data validation errors found:" in error_message:
                st.error("❌ **Excel File Validation Errors**")
                st.markdown("The following data validation errors were found in your Excel file:")
                
                # Split the error message to extract the validation details
                parts = error_message.split("Data validation errors found:")
                if len(parts) > 1:
                    validation_details = parts[1].strip()
                    # Display each validation error in an expandable section
                    with st.expander(" **Detailed Error Information**", expanded=True):
                        st.markdown(validation_details)
                
                st.markdown("---")
                st.markdown("###  **How to Fix:**")
                st.markdown("1. Open your Excel file")
                st.markdown("2. Navigate to the specific cells mentioned above")
                st.markdown("3. Ensure all numeric fields contain valid numbers (not letters or text)")
                st.markdown("4. Save the file and try uploading again")
                
                st.info(" **Tip:** The most common issue is entering letters in numeric fields like 'Testing and Commissioning' prices.")
                
            else:
                st.error(f"❌ Error reading Excel file: {error_message}")
                # Show detailed traceback for debugging
                with st.expander("🔍 Technical Details", expanded=False):
                    import traceback
                    st.code(traceback.format_exc())
            
            if os.path.exists(temp_path):
                os.remove(temp_path)

def initialize_session_state():
    """Initialize session state variables."""
    if 'uploaded_project_data' not in st.session_state:
        st.session_state.uploaded_project_data = None
    if 'upload_success' not in st.session_state:
        st.session_state.upload_success = False
    if 'levels' not in st.session_state:
        st.session_state.levels = []
    if 'current_step' not in st.session_state:
        st.session_state.current_step = 1
    if 'project_info' not in st.session_state:
        st.session_state.project_info = {}
    # Initialize template selection with default 19.2
    if "selected_template" not in st.session_state:
        st.session_state.selected_template = "Cost Sheet R19.2 Jun 2025"
    if "template_path" not in st.session_state:
        st.session_state.template_path = "templates/excel/Cost Sheet R19.2 Jun 2025.xlsx"

def navigation_buttons():
    """Display navigation buttons based on the current step."""
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col1:
        if st.session_state.current_step > 1:
            if st.button("← Previous", key="nav_prev"):
                st.session_state.current_step -= 1
                st.rerun()
    
    with col2:
        # Progress indicator
        steps = ["Project Info", "Structure", "Canopies", "Review"]
        current_step = st.session_state.current_step
        progress_text = f"Step {current_step}/4: {steps[current_step-1]}"
        st.markdown(f"<div style='text-align: center; font-weight: bold;'>{progress_text}</div>", unsafe_allow_html=True)
        
        # Progress bar
        progress = current_step / 4
        st.progress(progress)
    
    with col3:
        # Comment out navigation to step 4 for now
        if st.session_state.current_step < 3:  # Changed from 4 to 3
            if st.button("Next →", key="nav_next"):
                st.session_state.current_step += 1
                st.rerun()

def step1_project_information():
    """Step 1: Project Information"""
    st.header("Step 1: Project Information")
    
    # Use uploaded data if available
    if st.session_state.uploaded_project_data:
        project_data = st.session_state.uploaded_project_data.copy()
        if not st.session_state.project_info:  # Only show message once
            st.info("Form auto-populated from uploaded Excel file. You can modify any fields as needed.")
    else:
        project_data = st.session_state.project_info
    
    # Company selection mode (outside columns for immediate reactivity)
    company_mode = st.radio(
        "Company Selection *",
        options=["Select from list", "Enter custom company"],
        index=1 if project_data.get("company_mode", "Enter custom company") == "Enter custom company" else 0,
        key="company_mode_input",
        help="Choose whether to select from predefined companies or enter a custom company"
    )
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Initialize session state for project fields
        if "project_name_state" not in st.session_state:
            st.session_state.project_name_state = project_data.get('project_name', '')
        if "customer_state" not in st.session_state:
            st.session_state.customer_state = project_data.get('customer', '')
        
        def update_project_name():
            st.session_state.project_name_state = st.session_state.project_name
        
        def update_customer():
            st.session_state.customer_state = st.session_state.customer
            
        project_name = st.text_input("Project Name", 
                                   value=st.session_state.project_name_state, 
                                   key="project_name",
                                   on_change=update_project_name)
        customer = st.text_input("Customer Name", 
                                value=st.session_state.customer_state, 
                                key="customer",
                                on_change=update_customer)
        
        # Company and address selection based on mode
        if company_mode == "Select from list":
            # Get current company value and find its index
            current_company = project_data.get('company', '')
            company_options = list(COMPANY_ADDRESSES.keys())
            default_index = 0
            if current_company in company_options:
                default_index = company_options.index(current_company)
            
            company = st.selectbox(
                "Company *",
                options=company_options,
                index=default_index,
                key="company_select",
                help="Select the company from the predefined list"
            )
            
            # Auto-populate address based on selected company
            if company in COMPANY_ADDRESSES:
                address = st.text_area("Address", value=COMPANY_ADDRESSES[company], key="address", help="Address auto-populated from company selection")
            else:
                address = st.text_area("Address", value=project_data.get('address', ''), key="address")
            
            custom_company_name = ""
            custom_company_address = ""
        else:
            # Custom company mode
            company = ""
            
            # Initialize session state for custom company fields
            if "custom_company_name_state" not in st.session_state:
                st.session_state.custom_company_name_state = project_data.get('custom_company_name', project_data.get('company', ''))
            if "custom_company_address_state" not in st.session_state:
                st.session_state.custom_company_address_state = project_data.get('custom_company_address', project_data.get('address', ''))
            
            def update_custom_company_name():
                st.session_state.custom_company_name_state = st.session_state.custom_company_name_input
            
            def update_custom_company_address():
                st.session_state.custom_company_address_state = st.session_state.custom_company_address_input
            
            custom_company_name = st.text_input(
                "Custom Company Name *",
                value=st.session_state.custom_company_name_state,
                key="custom_company_name_input",
                help="Enter the custom company name",
                on_change=update_custom_company_name
            )
            custom_company_address = st.text_area(
                "Custom Company Address *",
                value=st.session_state.custom_company_address_state,
                key="custom_company_address_input",
                help="Enter the full company address (use line breaks for multiple lines)",
                height=100,
                on_change=update_custom_company_address
            )
            address = custom_company_address
        
        # Initialize session state for location
        if "location_state" not in st.session_state:
            st.session_state.location_state = project_data.get('project_location', '')
        
        def update_location():
            st.session_state.location_state = st.session_state.project_location
        
        location = st.text_input("Location", 
                                value=st.session_state.location_state, 
                                key="project_location",
                                on_change=update_location)
    
    with col2:
        # Initialize session state for project number
        if "project_number_state" not in st.session_state:
            st.session_state.project_number_state = project_data.get('project_number', '')
        
        def update_project_number():
            st.session_state.project_number_state = st.session_state.project_number
        
        project_number = st.text_input("Project Number", 
                                     value=st.session_state.project_number_state, 
                                     key="project_number",
                                     on_change=update_project_number)
        
        # Handle date conversion from string format
        if project_data.get('date'):
            try:
                if isinstance(project_data['date'], str):
                    # Try to parse date string in DD/MM/YYYY format
                    date_obj = datetime.strptime(project_data['date'], "%d/%m/%Y").date()
                    date = st.date_input("Date", value=date_obj, key="date")
                else:
                    date = st.date_input("Date", value=project_data['date'], key="date")
            except:
                date = st.date_input("Date", key="date")
        else:
            date = st.date_input("Date", key="date")
        
        # Get estimator options and set default from uploaded data
        estimator_options = list(ESTIMATORS.keys())
        default_estimator_index = 0
        if project_data.get('estimator'):
            try:
                default_estimator_index = estimator_options.index(project_data['estimator'])
            except ValueError:
                # If exact match not found, try partial match
                for i, estimator in enumerate(estimator_options):
                    if estimator.lower() in project_data['estimator'].lower() or project_data['estimator'].lower() in estimator.lower():
                        default_estimator_index = i
                        break
        
        estimator = st.selectbox("Estimator", estimator_options, index=default_estimator_index, key="estimator")
        
        # Get sales contact options and set default from uploaded data
        sales_contact_options = list(SALES_CONTACTS.keys())
        default_sales_contact_index = 0
        if project_data.get('sales_contact'):
            try:
                default_sales_contact_index = sales_contact_options.index(project_data['sales_contact'])
            except ValueError:
                # If exact match not found, try partial match
                for i, contact in enumerate(sales_contact_options):
                    if contact.lower() in project_data['sales_contact'].lower() or project_data['sales_contact'].lower() in contact.lower():
                        default_sales_contact_index = i
                        break
        
        sales_contact = st.selectbox("Sales Contact", sales_contact_options, index=default_sales_contact_index, key="sales_contact")
        
        # Get delivery location options and set default from uploaded data
        delivery_options = DELIVERY_LOCATIONS
        default_delivery_index = 0
        if project_data.get('delivery_location'):
            try:
                default_delivery_index = delivery_options.index(project_data['delivery_location'])
            except ValueError:
                # If exact match not found, keep default as 0 (Select...)
                pass
        
        delivery_location = st.selectbox("Delivery Location", delivery_options, index=default_delivery_index, key="delivery_location")
        
        # Initialize session state for revision
        if "revision_state" not in st.session_state:
            st.session_state.revision_state = project_data.get('revision', '')
        
        def update_revision():
            st.session_state.revision_state = st.session_state.revision
        
        # Revision field with uploaded data
        revision = st.text_input("Revision (leave blank for initial version)", 
                                value=st.session_state.revision_state, 
                                key="revision",
                                on_change=update_revision)
    
    # Project-level options
    st.markdown("---")
    st.subheader("Project Options")
    
    # Initialize contract option in session state
    if "contract_option_state" not in st.session_state:
        st.session_state.contract_option_state = project_data.get('contract_option', False)
    
    def update_contract_option():
        st.session_state.contract_option_state = st.session_state.contract_option
    
    contract_option = st.checkbox(
        "Include Contract Sheets",
        value=st.session_state.contract_option_state,
        key="contract_option",
        help="Include Contract, Spiral Duct, Supply Duct, and Extract Duct tabs in the Excel file",
        on_change=update_contract_option
    )
    
    # Determine final company name and address based on mode
    if company_mode == "Select from list":
        final_company_name = company
        final_address = COMPANY_ADDRESSES.get(company, address)
    else:  # Custom company
        final_company_name = st.session_state.get('custom_company_name_state', '')
        final_address = st.session_state.get('custom_company_address_state', '')
    
    # Note: uploaded_project_data is now used only for display purposes
    # All actual data is stored directly in session state variables
    
    # Store project info in session state using session state values for immediate updates
    st.session_state.project_info = {
        'project_name': st.session_state.project_name_state,
        'customer': st.session_state.customer_state,
        'company': final_company_name,
        'address': final_address,
        'project_location': st.session_state.location_state,
        'project_number': st.session_state.project_number_state,
        'date': date.strftime("%d/%m/%Y") if date else "",
        'estimator': estimator,
        'sales_contact': sales_contact,
        'delivery_location': delivery_location if delivery_location != "Select..." else "",
        'revision': st.session_state.revision_state,
        # Store the selection mode and custom fields for form persistence
        'company_mode': company_mode,
        'custom_company_name': st.session_state.get('custom_company_name_state', ''),
        'custom_company_address': st.session_state.get('custom_company_address_state', ''),
        # Project-level options
        'contract_option': st.session_state.contract_option_state
    }
    
    # Validation
    if company_mode == "Select from list":
        required_fields = [st.session_state.project_name_state, st.session_state.project_number_state, final_company_name]
    else:
        required_fields = [st.session_state.project_name_state, st.session_state.project_number_state, final_company_name, final_address]
    
    if all(required_fields):
        st.success("Project information is complete!")
    else:
        st.warning("Please fill in all required fields to continue.")

def step2_project_structure():
    """Step 2: Project Structure (Levels and Areas)"""
    st.header("Step 2: Project Structure")
    st.markdown("Define the levels and areas for your project.")
    
    # Note: Levels data is now loaded immediately when Excel is uploaded
    # No need to load it here as it's already in st.session_state.levels
    
    # Level management
    col1, col2 = st.columns([3, 1])
    with col1:
        st.subheader("Levels")
    with col2:
        if st.button("Add Level", key="add_level"):
            new_level_number = len(st.session_state.levels) + 1
            st.session_state.levels.append({
                "level_number": new_level_number,
                "level_name": f"Level {new_level_number}",
                "areas": []
            })
            st.rerun()

    # Display levels
    for level_idx, level in enumerate(st.session_state.levels):
        with st.expander(f"Level {level['level_number']}: {level['level_name']}", expanded=True):
            # Level name input with immediate update
            def update_level_name():
                try:
                    # Check if index is still valid before updating
                    if level_idx < len(st.session_state.levels):
                        st.session_state.levels[level_idx]['level_name'] = st.session_state[f"level_name_{level_idx}"]
                except (IndexError, KeyError) as e:
                    # Silently fail - the UI will rerender with correct indices
                    pass
            
            new_level_name = st.text_input(f"Level Name", 
                                         value=level['level_name'], 
                                         key=f"level_name_{level_idx}",
                                         on_change=update_level_name)
            
            # Remove level button
            if st.button(f"Remove Level {level['level_number']}", key=f"remove_level_{level_idx}"):
                del st.session_state.levels[level_idx]
                # Renumber remaining levels
                for i, remaining_level in enumerate(st.session_state.levels):
                    remaining_level['level_number'] = i + 1
                    remaining_level['level_name'] = remaining_level['level_name'].replace(f"Level {remaining_level['level_number']}", f"Level {i + 1}")
                st.rerun()
            
            # Area management for this level
            st.markdown(f"### Areas in {level['level_name']}")
            col1, col2 = st.columns([3, 1])
            with col2:
                if st.button(f"Add Area", key=f"add_area_{level_idx}"):
                    st.session_state.levels[level_idx]['areas'].append({
                        "name": f"Area {len(level['areas']) + 1}",
                        "canopies": [],
                        "options": {
                            "uvc": False, 
                            "recoair": False, 
                            "marvel": False, 
                            "uv_extra_over": False,
                            "vent_clg": False
                        }
                    })
                    st.rerun()
            
            # Display areas
            for area_idx, area in enumerate(level['areas']):
                area_key = f"level_{level_idx}_area_{area_idx}"
                with st.container():
                    st.markdown(f"#### Area: {area['name']}")
                    
                    # Area name and options
                    col1, col2, col3 = st.columns([2, 1, 1])
                    with col1:
                        # Initialize session state for area name if not exists
                        area_name_key = f"{area_key}_name_state"
                        if area_name_key not in st.session_state:
                            st.session_state[area_name_key] = area['name']
                        
                        def update_area_name():
                            try:
                                # Check if indices are still valid before updating
                                if (level_idx < len(st.session_state.levels) and 
                                    area_idx < len(st.session_state.levels[level_idx]['areas'])):
                                    st.session_state.levels[level_idx]['areas'][area_idx]['name'] = st.session_state[f"{area_key}_name"]
                                    st.session_state[area_name_key] = st.session_state[f"{area_key}_name"]
                            except (IndexError, KeyError) as e:
                                # If there's an error, just update the state key
                                st.session_state[area_name_key] = st.session_state.get(f"{area_key}_name", "")
                        
                        new_area_name = st.text_input("Area Name", 
                                                    value=st.session_state[area_name_key], 
                                                    key=f"{area_key}_name",
                                                    on_change=update_area_name)
                    
                    with col2:
                        # Area options
                        st.markdown("**Options:**")
                        
                        # Initialize session state for options if not exists
                        def update_area_options():
                            try:
                                # Check if indices are still valid before updating
                                if (level_idx < len(st.session_state.levels) and 
                                    area_idx < len(st.session_state.levels[level_idx]['areas'])):
                                    st.session_state.levels[level_idx]['areas'][area_idx]['options'] = {
                                        'uvc': st.session_state.get(f"{area_key}_uvc", False),
                                        'recoair': st.session_state.get(f"{area_key}_recoair", False),
                                        'marvel': st.session_state.get(f"{area_key}_marvel", False),
                                        'uv_extra_over': st.session_state.get(f"{area_key}_uv_extra_over", False),
                                        'vent_clg': st.session_state.get(f"{area_key}_vent_clg", False)
                                    }
                            except (IndexError, KeyError) as e:
                                # If there's an error, silently fail - the checkboxes will maintain their state
                                pass
                        
                        uvc = st.checkbox("UV-C", 
                                        value=area['options'].get('uvc', False), 
                                        key=f"{area_key}_uvc",
                                        on_change=update_area_options)
                        recoair = st.checkbox("RecoAir", 
                                            value=area['options'].get('recoair', False), 
                                            key=f"{area_key}_recoair",
                                            on_change=update_area_options)
                        
                        marvel = st.checkbox("Marvel", 
                                        value=area['options'].get('marvel', False), 
                                        key=f"{area_key}_marvel",
                                        on_change=update_area_options)
                        
                        # UV Extra Over option - always available regardless of canopies
                        uv_extra_over = st.checkbox("UV Extra Over", 
                                                  value=area['options'].get('uv_extra_over', False), 
                                                  key=f"{area_key}_uv_extra_over", 
                                                  help="Calculate additional cost for UV functionality",
                                                  on_change=update_area_options)
                        
                        vent_clg = st.checkbox("VENT CLG", 
                                            value=area['options'].get('vent_clg', False), 
                                            key=f"{area_key}_vent_clg",
                                            help="Toggle if Ventilated Ceiling is needed for this area",
                                            on_change=update_area_options)
                        
                        # Options are updated via the callback, no need for direct update here
                    
                    with col3:
                        if st.button(f"Remove Area", key=f"{area_key}_remove"):
                            del st.session_state.levels[level_idx]['areas'][area_idx]
                            st.rerun()
                    
                    st.markdown("---")

def step3_canopy_configuration():
    """Step 3: Canopy Configuration"""
    st.header("Step 3: Canopy Configuration")
    st.markdown("Configure canopies for each area.")
    
    if not st.session_state.levels:
        st.warning("Please add levels and areas in Step 2 before configuring canopies.")
        return
    
    # Display areas with canopy configuration
    for level_idx, level in enumerate(st.session_state.levels):
        st.subheader(f"Level {level['level_number']}: {level['level_name']}")
        
        for area_idx, area in enumerate(level['areas']):
            area_key = f"level_{level_idx}_area_{area_idx}"
            with st.expander(f"Area: {area['name']}", expanded=True):
                # Check if area has any UV canopies to determine available options
                has_uv_canopies = any(canopy.get('model', '').upper().startswith('UV') for canopy in area.get('canopies', []))
                
                # Display area options with UV Extra Over
                options_text = f"UV-C: {'Yes' if area['options']['uvc'] else 'No'} | RecoAir: {'Yes' if area['options']['recoair'] else 'No'} | Marvel: {'Yes' if area['options'].get('marvel', False) else 'No'}"
                options_text += f" | UV Extra Over: {'Yes' if area['options'].get('uv_extra_over', False) else 'No'}"
                
                st.markdown(f"**Area Options:** {options_text}")
                
                # Canopy management
                st.markdown("**Canopies:**")
                col1, col2 = st.columns([3, 1])
                with col2:
                    if st.button("Add Canopy", key=f"{area_key}_add_canopy"):
                        new_canopy = {
                            "reference_number": f"C{len(area['canopies']) + 1:03d}",
                            "configuration": "",
                            "model": "",
                            "length": 0,
                            "width": 0,
                            "height": 555,  # Default height set to 555
                            "sections": 0,
                            "lighting_type": "",
                            "extract_volume": "",
                            "extract_static": "",
                            "mua_volume": "",
                            "supply_static": "",
                            "sdu_item_number": "",
                            "options": {"fire_suppression": False, "sdu": False},
                            "wall_cladding": {"type": "None", "width": None, "height": None, "position": None}
                        }
                        st.session_state.levels[level_idx]['areas'][area_idx]['canopies'].append(new_canopy)
                        st.rerun()
                
                # Display canopies
                for canopy_idx, canopy in enumerate(area['canopies']):
                    canopy_key = f"{area_key}_canopy_{canopy_idx}"
                    
                    # Initialize session state for canopy fields if not already present
                    if f"{canopy_key}_ref" not in st.session_state:
                        st.session_state[f"{canopy_key}_ref"] = canopy.get('reference_number', '')
                    if f"{canopy_key}_model" not in st.session_state:
                        st.session_state[f"{canopy_key}_model"] = canopy.get('model', '')
                    if f"{canopy_key}_config" not in st.session_state:
                        st.session_state[f"{canopy_key}_config"] = canopy.get('configuration', '')
                    if f"{canopy_key}_length" not in st.session_state:
                        length_val = canopy.get('length', 0)
                        st.session_state[f"{canopy_key}_length"] = int(length_val) if length_val and str(length_val).strip() else 0
                    if f"{canopy_key}_width" not in st.session_state:
                        width_val = canopy.get('width', 0)
                        st.session_state[f"{canopy_key}_width"] = int(width_val) if width_val and str(width_val).strip() else 0
                    if f"{canopy_key}_height" not in st.session_state:
                        height_val = canopy.get('height', 555)
                        st.session_state[f"{canopy_key}_height"] = int(height_val) if height_val and str(height_val).strip() else 555
                    if f"{canopy_key}_sections" not in st.session_state:
                        sections_val = canopy.get('sections', 0)
                        st.session_state[f"{canopy_key}_sections"] = int(sections_val) if sections_val and str(sections_val).strip() else 0
                    if f"{canopy_key}_fire" not in st.session_state:
                        st.session_state[f"{canopy_key}_fire"] = canopy.get('options', {}).get('fire_suppression', False)
                    if f"{canopy_key}_sdu" not in st.session_state:
                        st.session_state[f"{canopy_key}_sdu"] = canopy.get('options', {}).get('sdu', False)
                    if f"{canopy_key}_sdu_item" not in st.session_state:
                        st.session_state[f"{canopy_key}_sdu_item"] = canopy.get('sdu_item_number', '')
                    
                    with st.container():
                        st.markdown(f"**Canopy {canopy_idx + 1}:**")
                        
                        # Basic canopy info - clean organized layout
                        
                        # Define update function for canopy data
                        def update_canopy_data():
                            try:
                                # Check if indices are still valid before updating
                                if (level_idx < len(st.session_state.levels) and 
                                    area_idx < len(st.session_state.levels[level_idx]['areas']) and 
                                    canopy_idx < len(st.session_state.levels[level_idx]['areas'][area_idx]['canopies'])):
                                    
                                    # Get the current canopy to preserve fields that don't have UI widgets
                                    current_canopy = st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]
                                    
                                    # Update only the fields that have UI widgets, preserving all other fields
                                    current_canopy.update({
                                        'reference_number': st.session_state.get(f"{canopy_key}_ref", ''),
                                        'model': st.session_state.get(f"{canopy_key}_model", ''),
                                        'configuration': st.session_state.get(f"{canopy_key}_config", ''),
                                        'length': st.session_state.get(f"{canopy_key}_length", 0),
                                        'width': st.session_state.get(f"{canopy_key}_width", 0),
                                        'height': st.session_state.get(f"{canopy_key}_height", 0),
                                        'sections': st.session_state.get(f"{canopy_key}_sections", 0),
                                        'sdu_item_number': st.session_state.get(f"{canopy_key}_sdu_item", ''),
                                        'options': {
                                            'fire_suppression': st.session_state.get(f"{canopy_key}_fire", False),
                                            'sdu': st.session_state.get(f"{canopy_key}_sdu", False)
                                        }
                                    })
                            except (IndexError, KeyError) as e:
                                # Silently ignore index errors - the UI will rerender with correct indices
                                pass

                        # Row 1: Reference, Model, Configuration
                        row1_col1, row1_col2, row1_col3 = st.columns(3)
                        
                        with row1_col1:
                            ref_num = st.text_input("Reference", 
                                                   key=f"{canopy_key}_ref")
                        
                        with row1_col2:
                            model_options = [""] + VALID_CANOPY_MODELS
                            model_index = 0
                            if canopy.get('model', '') in model_options:
                                model_index = model_options.index(canopy.get('model', ''))
                            
                            model = st.selectbox("Model", model_options,
                                               key=f"{canopy_key}_model")
                        
                        with row1_col3:
                            config_options = ["Wall", "Island"]
                            config_index = 0
                            if canopy.get('configuration', '') in config_options:
                                config_index = config_options.index(canopy.get('configuration', ''))
                            
                            configuration = st.selectbox("Configuration", config_options,
                                                       key=f"{canopy_key}_config")
                        
                        # Row 2: Dimensions - Length, Width, Height
                        st.markdown("**Dimensions:**")
                        row2_col1, row2_col2, row2_col3 = st.columns(3)
                        
                        with row2_col1:
                            length = st.number_input("Length", 
                                                   key=f"{canopy_key}_length",
                                                   min_value=0)
                        
                        with row2_col2:
                            width = st.number_input("Width", 
                                                  key=f"{canopy_key}_width",
                                                  min_value=0)
                        
                        with row2_col3:
                            # Use 555 as default height if no height is set
                            default_height = canopy.get('height', 555)
                            if default_height == 0 or default_height == "":
                                default_height = 555
                            height = st.number_input("Height", 
                                                   key=f"{canopy_key}_height",
                                                   min_value=0)
                        
                        # Row 3: Sections and Fire Suppression
                        row3_col1, row3_col2, row3_col3 = st.columns(3)
                        
                        with row3_col1:
                            sections = st.number_input("Sections", 
                                                     key=f"{canopy_key}_sections",
                                                     min_value=0)
                        
                        with row3_col2:
                            fire_suppression = st.checkbox("Fire Suppression", 
                                                          key=f"{canopy_key}_fire")
                        
                        with row3_col3:
                            sdu = st.checkbox("SDU", 
                                            key=f"{canopy_key}_sdu")
                        
                        # SDU Item Number input (only show if SDU is checked)
                        if st.session_state.get(f"{canopy_key}_sdu", False):
                            sdu_item_number = st.text_input(
                                "SDU Item Number",
                                key=f"{canopy_key}_sdu_item",
                                help="Enter the item number for this SDU (will be written to B12)"
                            )
                        
                        # Wall Cladding Section  
                        st.markdown("**Wall Cladding:**")
                        
                        # Initialize wall cladding state if not already present
                        if f"{canopy_key}_wall_cladding_enabled" not in st.session_state:
                            st.session_state[f"{canopy_key}_wall_cladding_enabled"] = canopy.get('wall_cladding', {}).get('type') not in ['None', None, '']
                        
                        wall_cladding_enabled = st.checkbox("With Wall Cladding", 
                                                          key=f"{canopy_key}_wall_cladding_enabled")
                        
                        if wall_cladding_enabled:
                            clad_col1, clad_col2, clad_col3 = st.columns(3)
                            
                            # Initialize wall cladding dimensions if not already present
                            if f"{canopy_key}_clad_width" not in st.session_state:
                                width_val = canopy.get('wall_cladding', {}).get('width', 0)
                                st.session_state[f"{canopy_key}_clad_width"] = int(width_val) if width_val and str(width_val).strip() else 0
                            if f"{canopy_key}_clad_height" not in st.session_state:
                                height_val = canopy.get('wall_cladding', {}).get('height', 0)
                                st.session_state[f"{canopy_key}_clad_height"] = int(height_val) if height_val and str(height_val).strip() else 0
                            
                            with clad_col1:
                                cladding_width = st.number_input(
                                    "Width (mm)", 
                                    key=f"{canopy_key}_clad_width",
                                    min_value=0
                                )
                            
                            with clad_col2:
                                cladding_height = st.number_input(
                                    "Height (mm)", 
                                    key=f"{canopy_key}_clad_height",
                                    min_value=0
                                )
                            
                            with clad_col3:
                                # Initialize position if not already present
                                if f"{canopy_key}_clad_position" not in st.session_state:
                                    current_positions = canopy.get('wall_cladding', {}).get('position', [])
                                    if isinstance(current_positions, str):
                                        current_positions = [current_positions] if current_positions else []
                                    elif current_positions is None:
                                        current_positions = []
                                    st.session_state[f"{canopy_key}_clad_position"] = current_positions
                                
                                cladding_positions = st.multiselect(
                                    "Position",
                                    options=["rear", "left hand", "right hand"],
                                    key=f"{canopy_key}_clad_position"
                                )
                        
                        # Canopy data is updated via callbacks
                        
                        # Remove canopy button
                        if st.button(f"Remove Canopy", key=f"{canopy_key}_remove"):
                            del st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]
                            st.rerun()
                        
                        st.markdown("---")
    
    # Add Excel generation section at the bottom of Step 3
    st.markdown("---")
    st.subheader("💾 Save Your Work")
    st.markdown("Generate an Excel file to save your current configuration. You can continue editing or proceed to the next step.")
    
    # Show Excel generation section
    generate_excel_section()

def generate_excel_section():
    """Reusable Excel generation section that can be used in multiple steps."""
    if st.button("Generate Excel Cost Sheet", type="primary", use_container_width=True):
        try:
            # Combine all project data
            final_project_data = st.session_state.project_info.copy()
            final_project_data['levels'] = st.session_state.levels
            
            # Generate Excel file using selected template
            template_path = st.session_state.get('template_path', 'templates/excel/Cost Sheet R19.1 May 2025.xlsx')
            with st.spinner("Generating Excel cost sheet..."):
                output_path = save_to_excel(final_project_data, template_path)
            
            st.success(f"Excel cost sheet generated successfully!")
            
            # Provide download option for Excel file
            try:
                with open(output_path, "rb") as file:
                    excel_data = file.read()
                
                # Create download filename
                project_number = final_project_data.get('project_number', 'unknown')
                date_str = final_project_data.get('date', '')
                if date_str:
                    formatted_date = date_str.replace('/', '')
                else:
                    formatted_date = get_current_date().replace('/', '')
                
                download_filename = f"{project_number} Cost Sheet {formatted_date}.xlsx"
                
                st.download_button(
                    label="Download Excel Cost Sheet",
                    data=excel_data,
                    file_name=download_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    type="primary"
                )
                
                # Clean up
                if os.path.exists(output_path):
                    os.remove(output_path)
                    
            except Exception as e:
                st.error(f"Error preparing download: {str(e)}")
                
        except Exception as e:
            st.error(f"Error generating Excel: {str(e)}")
            st.exception(e)

def step4_review_and_generate():
    """Step 4: Review and Generate"""
    st.header("Step 4: Review & Generate")
    st.markdown("Review your project configuration and generate the cost sheet.")
    
    # Validation
    if not st.session_state.project_info.get('project_name') or not st.session_state.project_info.get('project_number'):
        st.error("Please complete Step 1: Project Information")
        return
    
    # Project summary
    st.subheader("Project Summary")
    col1, col2 = st.columns(2)
    
    with col1:
        st.write("**Project Name:**", st.session_state.project_info.get('project_name'))
        st.write("**Customer:**", st.session_state.project_info.get('customer'))
        st.write("**Company:**", st.session_state.project_info.get('company'))
        st.write("**Address:**", st.session_state.project_info.get('address'))
        st.write("**Location:**", st.session_state.project_info.get('project_location'))
    
    with col2:
        st.write("**Project Number:**", st.session_state.project_info.get('project_number'))
        st.write("**Date:**", st.session_state.project_info.get('date'))
        st.write("**Estimator:**", st.session_state.project_info.get('estimator'))
        st.write("**Sales Contact:**", st.session_state.project_info.get('sales_contact'))
        st.write("**Delivery Location:**", st.session_state.project_info.get('delivery_location'))
        st.write("**Revision:**", st.session_state.project_info.get('revision') or 'Initial Version')
        st.write("**Contract Sheets:**", "Yes" if st.session_state.project_info.get('contract_option', False) else "No")
    
    # Structure summary
    st.subheader("Project Structure")
    total_levels = len(st.session_state.levels)
    total_areas = sum(len(level.get('areas', [])) for level in st.session_state.levels)
    total_canopies = sum(len(area.get('canopies', [])) for level in st.session_state.levels for area in level.get('areas', []))
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Levels", total_levels)
    with col2:
        st.metric("Areas", total_areas)
    with col3:
        st.metric("Canopies", total_canopies)
    
    # Detailed structure
    if st.session_state.levels:
        with st.expander("Detailed Structure", expanded=False):
            for level in st.session_state.levels:
                st.markdown(f"**{level['level_name']}**")
                for area in level['areas']:
                    canopy_count = len(area['canopies'])
                    options = []
                    if area['options']['uvc']: options.append("UV-C")
                    if area['options']['recoair']: options.append("RecoAir")
                    if area['options']['marvel']: options.append("Marvel")
                    options_str = ", ".join(options) if options else "None"
                    st.write(f"  • {area['name']}: {canopy_count} canopies, Options: {options_str}")
    
    # Generate button
    st.markdown("---")
    generate_excel_section()

def populate_session_state_from_uploaded_data(extracted_data):
    """
    Populate all session state variables with data from uploaded Excel file.
    This ensures that all steps (project info, structure, canopies) are pre-filled.
    """
    try:
        # Clear any existing session state for form fields to force update
        form_fields_to_clear = [
            'project_name_state', 'customer_state', 'location_state', 
            'project_number_state', 'revision_state', 'custom_company_name_state', 
            'custom_company_address_state'
        ]
        for field in form_fields_to_clear:
            if field in st.session_state:
                del st.session_state[field]
        
        # Determine company mode based on whether company is in predefined list
        from config.business_data import COMPANY_ADDRESSES
        company_name = extracted_data.get('company', '')
        is_predefined_company = company_name in COMPANY_ADDRESSES.keys()
        
        # Populate project information
        st.session_state.project_info = {
            'project_name': extracted_data.get('project_name', ''),
            'customer': extracted_data.get('customer', ''),
            'company': company_name,
            'address': extracted_data.get('address', ''),
            'project_location': extracted_data.get('project_location', ''),
            'project_number': extracted_data.get('project_number', ''),
            'date': extracted_data.get('date', ''),
            'estimator': extracted_data.get('estimator', ''),
            'sales_contact': extracted_data.get('sales_contact', ''),
            'delivery_location': extracted_data.get('delivery_location', ''),
            'revision': extracted_data.get('revision', ''),
            'company_mode': 'Enter custom company' if not is_predefined_company else 'Select from list',
            'custom_company_name': company_name if not is_predefined_company else '',
            'custom_company_address': extracted_data.get('address', '') if not is_predefined_company else '',
            'contract_option': extracted_data.get('contract_option', False)
        }
        
        # Populate levels and areas structure
        if extracted_data.get('levels'):
            st.session_state.levels = extracted_data['levels'].copy()
        
        # Store template information if available in the extracted data
        if extracted_data.get('template_used'):
            # Map short template versions to full template names
            template_version_mapping = {
                'R19.2': "Cost Sheet R19.2 Jun 2025",
                'R19.1': "Cost Sheet R19.1 May 2025", 
                'R18.1': "Cost Sheet R18.1 (Legacy)",
                # Also handle full names in case they're already correct
                "Cost Sheet R19.2 Jun 2025": "Cost Sheet R19.2 Jun 2025",
                "Cost Sheet R19.1 May 2025": "Cost Sheet R19.1 May 2025",
                "Cost Sheet R18.1 (Legacy)": "Cost Sheet R18.1 (Legacy)"
            }
            
            template_options = {
                "Cost Sheet R19.2 Jun 2025": "templates/excel/Cost Sheet R19.2 Jun 2025.xlsx",
                "Cost Sheet R19.1 May 2025": "templates/excel/Cost Sheet R19.1 May 2025.xlsx",
                "Cost Sheet R18.1 (Legacy)": "templates/excel/Halton Cost Sheet Jan 2025.xlsx"
            }
            
            # Map the extracted template to the correct full name
            extracted_template = extracted_data['template_used']
            mapped_template = template_version_mapping.get(extracted_template, "Cost Sheet R19.2 Jun 2025")
            
            # Only set if the mapped template exists in current options
            if mapped_template in template_options:
                st.session_state.selected_template = mapped_template
                st.session_state.template_path = template_options[mapped_template]
                print(f"✅ Mapped template '{extracted_template}' to '{mapped_template}'")
            else:
                # Fallback to default
                st.session_state.selected_template = "Cost Sheet R19.2 Jun 2025"
                st.session_state.template_path = template_options["Cost Sheet R19.2 Jun 2025"]
                print(f"⚠️ Template '{extracted_template}' not recognized, using default")
        
        print(f"✅ Session state populated with uploaded data:")
        print(f"   - Project: {extracted_data.get('project_name', 'N/A')}")
        print(f"   - Levels: {len(extracted_data.get('levels', []))}")
        total_areas = sum(len(level.get('areas', [])) for level in extracted_data.get('levels', []))
        total_canopies = sum(len(area.get('canopies', [])) for level in extracted_data.get('levels', []) for area in level.get('areas', []))
        print(f"   - Areas: {total_areas}")
        print(f"   - Canopies: {total_canopies}")
        
    except Exception as e:
        print(f"❌ Error populating session state from uploaded data: {str(e)}")
        st.error(f"Error populating form data: {str(e)}")

def single_page_project_builder():
    """Single page project setup with sidebar structure builder."""
    st.title("🏗️ Single Page Project Builder")
    st.markdown("Build your project structure in the sidebar and configure canopies here.")
    
    # Load state from URL if present
    loaded_from_url = load_from_url()
    if loaded_from_url:
        st.success("✅ Progress loaded from saved link!")
    
    # Initialize session state if needed
    if 'project_info' not in st.session_state:
        st.session_state.project_info = {}
    if 'levels' not in st.session_state:
        st.session_state.levels = []
    if 'template_path' not in st.session_state:
        st.session_state.template_path = 'templates/excel/Cost Sheet R19.2 Jun 2025.xlsx'
    
    # Add save progress button
    add_save_progress_button()
    
    # Sidebar Project Info Section
    with st.sidebar:
        # Excel Upload Section
        st.markdown("### 📤 Load Existing Project")
        with st.expander("Upload Excel (Optional)", expanded=False):
            uploaded_file = st.file_uploader(
                "Choose Excel file to pre-fill data",
                type=['xlsx', 'xls'],
                key="sp_excel_upload",
                help="Upload a previously generated cost sheet to pre-fill the form"
            )
            
            if uploaded_file is not None:
                # Create a unique key for this file to track if we've already loaded it
                file_key = f"{uploaded_file.name}_{uploaded_file.size}"
                
                # Check if we've already loaded this file
                if 'sp_loaded_file' not in st.session_state or st.session_state.sp_loaded_file != file_key:
                    try:
                        # Save uploaded file temporarily
                        temp_path = f"temp_upload_{uploaded_file.name}"
                        with open(temp_path, "wb") as f:
                            f.write(uploaded_file.getbuffer())
                        
                        # Read project data from Excel
                        with st.spinner("Reading project data..."):
                            extracted_data = read_excel_project_data(temp_path)
                        
                        # Populate session state
                        if extracted_data:
                            st.session_state.project_info = {
                                'project_name': extracted_data.get('project_name', ''),
                                'project_number': extracted_data.get('project_number', ''),
                                'customer': extracted_data.get('customer', ''),
                                'company': extracted_data.get('company', ''),
                                'address': extracted_data.get('address', ''),
                                'project_location': extracted_data.get('project_location', ''),
                                'delivery_location': extracted_data.get('delivery_location', ''),
                                'estimator': extracted_data.get('estimator', ''),
                                'sales_contact': extracted_data.get('sales_contact', ''),
                                'date': extracted_data.get('date', get_current_date()),
                                'company_mode': 'Enter custom company' if extracted_data.get('company') not in COMPANY_ADDRESSES else 'Select from list'
                            }
                            
                            if 'levels' in extracted_data:
                                st.session_state.levels = extracted_data['levels']
                            
                            # Mark this file as loaded
                            st.session_state.sp_loaded_file = file_key
                            
                            st.success("✅ Data loaded successfully!")
                            
                            # Clean up
                            os.remove(temp_path)
                        
                    except Exception as e:
                        st.error(f"Error loading file: {str(e)}")
                else:
                    st.info("📄 File already loaded. Remove and re-upload to reload.")
        
        st.markdown("---")
        
        # Template Selection
        st.markdown("### 📄 Template Selection")
        template_options = {
            "Cost Sheet R19.2 Jun 2025": "templates/excel/Cost Sheet R19.2 Jun 2025.xlsx",
            "Cost Sheet R19.1 May 2025": "templates/excel/Cost Sheet R19.1 May 2025.xlsx",
            "Cost Sheet R18.1 (Legacy)": "templates/excel/Halton Cost Sheet Jan 2025.xlsx"
        }
        
        selected_template = st.selectbox(
            "Select Excel Template",
            options=list(template_options.keys()),
            index=0,
            key="sp_template_select",
            help="Choose which version of the cost sheet template to use"
        )
        st.session_state.template_path = template_options[selected_template]
        
        st.markdown("---")
        
        # Clear All Button
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ Clear All", key="sp_clear_all", help="Clear all data and start fresh"):
                st.session_state.project_info = {}
                st.session_state.levels = []
                if 'sp_loaded_file' in st.session_state:
                    del st.session_state.sp_loaded_file
                st.rerun()
        
        with col2:
            if st.button("➕ Add Sample", key="sp_add_sample", help="Add sample data for testing"):
                st.session_state.project_info = {
                    'project_name': 'Sample Kitchen Project',
                    'project_number': 'DEMO-001',
                    'customer': 'Demo Customer',
                    'company': 'Halton Company Ltd',
                    'address': COMPANY_ADDRESSES['Halton Company Ltd'],
                    'project_location': 'London',
                    'delivery_location': 'LONDON in FORS GOLD(varies)',
                    'estimator': list(ESTIMATORS.keys())[0],
                    'sales_contact': list(SALES_CONTACTS.keys())[0],
                    'date': get_current_date(),
                    'company_mode': 'Select from list'
                }
                st.session_state.levels = [{
                    'level_number': 1,
                    'level_name': 'Ground Floor',
                    'areas': [{
                        'name': 'Main Kitchen',
                        'canopies': [],
                        'options': {
                            'uvc': True,
                            'recoair': False,
                            'marvel': False,
                            'uv_extra_over': False,
                            'vent_clg': False
                        }
                    }]
                }]
                st.rerun()
        
        st.markdown("---")
        st.markdown("### 📋 Project Information")
        with st.expander("Project Details", expanded=True):
            # Project Name
            project_name = st.text_input(
                "Project Name",
                value=st.session_state.project_info.get('project_name', ''),
                key="sp_project_name"
            )
            if project_name:
                st.session_state.project_info['project_name'] = project_name
            
            # Customer
            customer = st.text_input(
                "Customer Name",
                value=st.session_state.project_info.get('customer', ''),
                key="sp_customer"
            )
            if customer:
                st.session_state.project_info['customer'] = customer
            
            # Company selection mode
            company_mode = st.radio(
                "Company Selection *",
                options=["Select from list", "Enter custom company"],
                index=1 if st.session_state.project_info.get("company_mode", "Enter custom company") == "Enter custom company" else 0,
                key="sp_company_mode",
                help="Choose whether to select from predefined companies or enter a custom company"
            )
            st.session_state.project_info['company_mode'] = company_mode
            
            # Company and address selection based on mode
            if company_mode == "Select from list":
                # Get current company value and find its index
                current_company = st.session_state.project_info.get('company', '')
                company_options = list(COMPANY_ADDRESSES.keys())
                default_index = 0
                if current_company in company_options:
                    default_index = company_options.index(current_company)
                
                company = st.selectbox(
                    "Company *",
                    options=company_options,
                    index=default_index,
                    key="sp_company_select",
                    help="Select the company from the predefined list"
                )
                st.session_state.project_info['company'] = company
                
                # Auto-populate address based on selected company
                if company in COMPANY_ADDRESSES:
                    address = st.text_area("Address", value=COMPANY_ADDRESSES[company], key="sp_address", disabled=True, help="Address auto-populated from company selection")
                    st.session_state.project_info['address'] = COMPANY_ADDRESSES[company]
                else:
                    address = st.text_area("Address", value=st.session_state.project_info.get('address', ''), key="sp_address")
                    st.session_state.project_info['address'] = address
            else:
                # Custom company mode
                custom_company_name = st.text_input(
                    "Custom Company Name *",
                    value=st.session_state.project_info.get('custom_company_name', st.session_state.project_info.get('company', '')),
                    key="sp_custom_company_name",
                    help="Enter the custom company name"
                )
                st.session_state.project_info['custom_company_name'] = custom_company_name
                st.session_state.project_info['company'] = custom_company_name
                
                custom_company_address = st.text_area(
                    "Custom Company Address *",
                    value=st.session_state.project_info.get('custom_company_address', st.session_state.project_info.get('address', '')),
                    key="sp_custom_company_address",
                    help="Enter the full company address (use line breaks for multiple lines)",
                    height=100
                )
                st.session_state.project_info['custom_company_address'] = custom_company_address
                st.session_state.project_info['address'] = custom_company_address
            
            # Project Number
            project_number = st.text_input(
                "Project Number",
                value=st.session_state.project_info.get('project_number', ''),
                key="sp_project_number"
            )
            if project_number:
                st.session_state.project_info['project_number'] = project_number
            
            # Date
            date_str = st.session_state.project_info.get('date', get_current_date())
            st.session_state.project_info['date'] = date_str
            st.text_input("Date", value=date_str, disabled=True)
            
            # Estimator
            estimator_options = list(ESTIMATORS.keys())
            current_estimator = st.session_state.project_info.get('estimator', '')
            default_estimator_index = 0
            if current_estimator in estimator_options:
                default_estimator_index = estimator_options.index(current_estimator)
            
            estimator = st.selectbox(
                "Estimator",
                estimator_options,
                index=default_estimator_index,
                key="sp_estimator"
            )
            st.session_state.project_info['estimator'] = estimator
            
            # Sales Contact
            sales_contact_options = list(SALES_CONTACTS.keys())
            current_sales_contact = st.session_state.project_info.get('sales_contact', '')
            default_sales_contact_index = 0
            if current_sales_contact in sales_contact_options:
                default_sales_contact_index = sales_contact_options.index(current_sales_contact)
            
            sales_contact = st.selectbox(
                "Sales Contact",
                sales_contact_options,
                index=default_sales_contact_index,
                key="sp_sales_contact"
            )
            st.session_state.project_info['sales_contact'] = sales_contact
            
            # Project Location
            project_location = st.text_input(
                "Project Location",
                value=st.session_state.project_info.get('project_location', ''),
                key="sp_project_location"
            )
            if project_location:
                st.session_state.project_info['project_location'] = project_location
            
            # Delivery Location
            delivery_options = DELIVERY_LOCATIONS
            current_delivery = st.session_state.project_info.get('delivery_location', 'Select...')
            default_delivery_index = 0
            if current_delivery in delivery_options:
                default_delivery_index = delivery_options.index(current_delivery)
            
            delivery_location = st.selectbox(
                "Delivery Location",
                options=delivery_options,
                index=default_delivery_index,
                key="sp_delivery_location"
            )
            st.session_state.project_info['delivery_location'] = delivery_location
            
            # Contract Option
            contract_option = st.checkbox(
                "Include Contract Sheets",
                value=st.session_state.project_info.get('contract_option', False),
                key="sp_contract_option",
                help="Include CONTRACT, EXTRACT DUCT, SUPPLY DUCT, and SPIRAL DUCT sheets in the Excel file"
            )
            st.session_state.project_info['contract_option'] = contract_option
        
        # Project Structure Section
        st.markdown("---")
        st.markdown("### 🏢 Project Structure")
        
        # Add Level button
        if st.button("➕ Add Level", key="sp_add_level", use_container_width=True):
            new_level_number = len(st.session_state.levels) + 1
            st.session_state.levels.append({
                "level_number": new_level_number,
                "level_name": f"Level {new_level_number}",
                "areas": []
            })
            st.rerun()
        
        # Display levels in sidebar
        for level_idx, level in enumerate(st.session_state.levels):
            st.markdown(f"#### {level['level_name']}")
            
            # Level controls
            col1, col2, col3 = st.columns([2, 1, 1])
            with col1:
                def update_sp_level_name(idx):
                    def callback():
                        st.session_state.levels[idx]['level_name'] = st.session_state[f"sp_level_name_{idx}"]
                    return callback
                
                new_level_name = st.text_input(
                    "Name",
                    value=level['level_name'],
                    key=f"sp_level_name_{level_idx}",
                    label_visibility="collapsed",
                    on_change=update_sp_level_name(level_idx)
                )
            
            with col2:
                if st.button("➕", key=f"sp_add_area_{level_idx}", help="Add Area"):
                    st.session_state.levels[level_idx]['areas'].append({
                        "name": f"Area {len(level['areas']) + 1}",
                        "canopies": [],
                        "options": {
                            "uvc": False,
                            "recoair": False,
                            "marvel": False,
                            "uv_extra_over": False,
                            "vent_clg": False
                        }
                    })
                    st.rerun()
            
            with col3:
                if st.button("🗑️", key=f"sp_del_level_{level_idx}", help="Delete Level"):
                    del st.session_state.levels[level_idx]
                    # Renumber remaining levels
                    for i, remaining_level in enumerate(st.session_state.levels):
                        remaining_level['level_number'] = i + 1
                    st.rerun()
            
            # Display areas for this level
            for area_idx, area in enumerate(level['areas']):
                with st.container():
                    # Area name and delete button
                    col1, col2 = st.columns([4, 1])
                    with col1:
                        def update_sp_area_name(l_idx, a_idx):
                            def callback():
                                st.session_state.levels[l_idx]['areas'][a_idx]['name'] = st.session_state[f"sp_area_name_{l_idx}_{a_idx}"]
                            return callback
                        
                        new_area_name = st.text_input(
                            "Area",
                            value=area['name'],
                            key=f"sp_area_name_{level_idx}_{area_idx}",
                            label_visibility="collapsed",
                            on_change=update_sp_area_name(level_idx, area_idx)
                        )
                    
                    with col2:
                        if st.button("🗑️", key=f"sp_del_area_{level_idx}_{area_idx}", help="Delete Area"):
                            del st.session_state.levels[level_idx]['areas'][area_idx]
                            st.rerun()
                    
                    # Area options with smaller columns
                    col1, col2 = st.columns(2)
                    with col1:
                        uvc = st.checkbox("UV-C", value=area['options'].get('uvc', False), 
                                         key=f"sp_uvc_{level_idx}_{area_idx}")
                        recoair = st.checkbox("RecoAir", value=area['options'].get('recoair', False),
                                            key=f"sp_recoair_{level_idx}_{area_idx}")
                        marvel = st.checkbox("Marvel", value=area['options'].get('marvel', False),
                                           key=f"sp_marvel_{level_idx}_{area_idx}")
                    
                    with col2:
                        uv_extra = st.checkbox("UV Extra", value=area['options'].get('uv_extra_over', False),
                                             key=f"sp_uv_extra_{level_idx}_{area_idx}")
                        vent_clg = st.checkbox("VENT CLG", value=area['options'].get('vent_clg', False),
                                             key=f"sp_vent_clg_{level_idx}_{area_idx}")
                    
                    # Update options
                    st.session_state.levels[level_idx]['areas'][area_idx]['options'] = {
                        'uvc': uvc,
                        'recoair': recoair,
                        'marvel': marvel,
                        'uv_extra_over': uv_extra,
                        'vent_clg': vent_clg
                    }
                    
                    st.markdown("---")
        
        # Excel generation at bottom of sidebar
        st.markdown("---")
        if st.button("💾 Generate Excel", key="sp_generate_excel", use_container_width=True, type="primary"):
            # Reuse the generate_excel_section logic
            try:
                final_project_data = st.session_state.project_info.copy()
                final_project_data['levels'] = st.session_state.levels
                
                template_path = st.session_state.get('template_path', 'templates/excel/Cost Sheet R19.1 May 2025.xlsx')
                with st.spinner("Generating Excel cost sheet..."):
                    output_path = save_to_excel(final_project_data, template_path)
                
                st.success("Excel generated!")
                
                # Provide download
                with open(output_path, "rb") as file:
                    excel_data = file.read()
                
                project_number = final_project_data.get('project_number', 'unknown')
                date_str = final_project_data.get('date', get_current_date()).replace('/', '')
                download_filename = f"{project_number} Cost Sheet {date_str}.xlsx"
                
                st.download_button(
                    label="📥 Download",
                    data=excel_data,
                    file_name=download_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="sp_download_excel"
                )
                
                if os.path.exists(output_path):
                    os.remove(output_path)
                    
            except Exception as e:
                st.error(f"Error: {str(e)}")
    
    # Main area - Canopy Configuration
    if not st.session_state.levels:
        st.info("👈 Start by adding levels and areas in the sidebar")
        return
    
    # Check if any areas exist
    total_areas = sum(len(level.get('areas', [])) for level in st.session_state.levels)
    if total_areas == 0:
        st.info("👈 Add areas to your levels in the sidebar")
        return
    
    st.markdown("## 🏗️ Canopy Configuration")
    st.markdown("Configure canopies for each area below:")
    
    # Display canopy configuration for each level and area
    for level_idx, level in enumerate(st.session_state.levels):
        if level['areas']:  # Only show if level has areas
            st.markdown(f"### {level['level_name']}")
            
            for area_idx, area in enumerate(level['areas']):
                with st.expander(f"**{area['name']}** - {len(area.get('canopies', []))} canopies", expanded=True):
                    # Show area options
                    options = []
                    if area['options'].get('uvc'): options.append("UV-C")
                    if area['options'].get('recoair'): options.append("RecoAir") 
                    if area['options'].get('marvel'): options.append("Marvel")
                    if area['options'].get('uv_extra_over'): options.append("UV Extra Over")
                    if area['options'].get('vent_clg'): options.append("VENT CLG")
                    
                    if options:
                        st.markdown(f"**Area Options:** {', '.join(options)}")
                    
                    # Add canopy button
                    if st.button(f"➕ Add Canopy", key=f"sp_add_canopy_{level_idx}_{area_idx}"):
                        new_canopy = {
                            "reference_number": f"C{len(area['canopies']) + 1:03d}",
                            "configuration": "",
                            "model": "",
                            "length": 0,
                            "width": 0,
                            "height": 555,
                            "sections": 0,
                            "lighting_type": "",
                            "extract_volume": "",
                            "extract_static": "",
                            "mua_volume": "",
                            "supply_static": "",
                            "sdu_item_number": "",
                            "options": {"fire_suppression": False, "sdu": False},
                            "wall_cladding": {"type": "None", "width": None, "height": None, "position": None}
                        }
                        st.session_state.levels[level_idx]['areas'][area_idx]['canopies'].append(new_canopy)
                        st.rerun()
                    
                    # Display existing canopies
                    for canopy_idx, canopy in enumerate(area.get('canopies', [])):
                        st.markdown(f"#### Canopy {canopy_idx + 1}")
                        
                        # Create a unique key prefix for this canopy
                        canopy_key = f"sp_canopy_{level_idx}_{area_idx}_{canopy_idx}"
                        
                        # Row 1: Basic info
                        col1, col2, col3, col4 = st.columns([2, 2, 2, 1])
                        
                        with col1:
                            ref_num = st.text_input(
                                "Reference",
                                value=canopy.get('reference_number', ''),
                                key=f"{canopy_key}_ref"
                            )
                            if ref_num != canopy.get('reference_number'):
                                st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['reference_number'] = ref_num
                        
                        with col2:
                            model_options = [""] + VALID_CANOPY_MODELS
                            model = st.selectbox(
                                "Model",
                                model_options,
                                index=model_options.index(canopy.get('model', '')) if canopy.get('model', '') in model_options else 0,
                                key=f"{canopy_key}_model"
                            )
                            if model != canopy.get('model'):
                                st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['model'] = model
                        
                        with col3:
                            config_options = ["Wall", "Island", "Single", "Double"]
                            config = st.selectbox(
                                "Configuration",
                                config_options,
                                index=config_options.index(canopy.get('configuration', 'Wall')) if canopy.get('configuration') in config_options else 0,
                                key=f"{canopy_key}_config"
                            )
                            if config != canopy.get('configuration'):
                                st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['configuration'] = config
                        
                        with col4:
                            if st.button("🗑️", key=f"{canopy_key}_delete", help="Delete Canopy"):
                                del st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]
                                st.rerun()
                        
                        # Row 2: Dimensions
                        col1, col2, col3, col4 = st.columns(4)
                        
                        with col1:
                            # Safe conversion to int
                            length_val = canopy.get('length', 0)
                            if length_val is None or length_val == '':
                                length_val = 0
                            try:
                                length_val = int(length_val)
                            except (ValueError, TypeError):
                                length_val = 0
                                
                            length = st.number_input(
                                "Length",
                                value=length_val,
                                key=f"{canopy_key}_length",
                                min_value=0
                            )
                            if length != canopy.get('length'):
                                st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['length'] = length
                        
                        with col2:
                            # Safe conversion to int
                            width_val = canopy.get('width', 0)
                            if width_val is None or width_val == '':
                                width_val = 0
                            try:
                                width_val = int(width_val)
                            except (ValueError, TypeError):
                                width_val = 0
                                
                            width = st.number_input(
                                "Width",
                                value=width_val,
                                key=f"{canopy_key}_width",
                                min_value=0
                            )
                            if width != canopy.get('width'):
                                st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['width'] = width
                        
                        with col3:
                            # Safe conversion to int
                            height_val = canopy.get('height', 555)
                            if height_val is None or height_val == '':
                                height_val = 555
                            try:
                                height_val = int(height_val)
                            except (ValueError, TypeError):
                                height_val = 555
                                
                            height = st.number_input(
                                "Height",
                                value=height_val,
                                key=f"{canopy_key}_height",
                                min_value=0
                            )
                            if height != canopy.get('height'):
                                st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['height'] = height
                        
                        with col4:
                            # Safe conversion to int
                            sections_val = canopy.get('sections', 0)
                            if sections_val is None or sections_val == '':
                                sections_val = 0
                            try:
                                sections_val = int(sections_val)
                            except (ValueError, TypeError):
                                sections_val = 0
                                
                            sections = st.number_input(
                                "Sections",
                                value=sections_val,
                                key=f"{canopy_key}_sections",
                                min_value=0
                            )
                            if sections != canopy.get('sections'):
                                st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['sections'] = sections
                        
                        # Row 3: Options
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            fire_supp = st.checkbox(
                                "Fire Suppression",
                                value=canopy.get('options', {}).get('fire_suppression', False),
                                key=f"{canopy_key}_fire"
                            )
                            if fire_supp != canopy.get('options', {}).get('fire_suppression'):
                                st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['options']['fire_suppression'] = fire_supp
                        
                        with col2:
                            sdu = st.checkbox(
                                "SDU",
                                value=canopy.get('options', {}).get('sdu', False),
                                key=f"{canopy_key}_sdu"
                            )
                            if sdu != canopy.get('options', {}).get('sdu'):
                                st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['options']['sdu'] = sdu
                        
                        with col3:
                            if sdu:
                                sdu_item = st.text_input(
                                    "SDU Item Number",
                                    value=canopy.get('sdu_item_number', ''),
                                    key=f"{canopy_key}_sdu_item"
                                )
                                if sdu_item != canopy.get('sdu_item_number'):
                                    st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['sdu_item_number'] = sdu_item
                        
                        # Wall Cladding Section
                        st.markdown("**Wall Cladding:**")
                        wall_clad = st.checkbox(
                            "With Wall Cladding",
                            value=canopy.get('wall_cladding', {}).get('type', 'None') != 'None',
                            key=f"{canopy_key}_wall_clad"
                        )
                        
                        if wall_clad:
                            clad_col1, clad_col2, clad_col3 = st.columns(3)
                            
                            with clad_col1:
                                # Safe conversion to int
                                clad_width_val = canopy.get('wall_cladding', {}).get('width', 0)
                                if clad_width_val is None or clad_width_val == '':
                                    clad_width_val = 0
                                try:
                                    clad_width_val = int(clad_width_val)
                                except (ValueError, TypeError):
                                    clad_width_val = 0
                                    
                                clad_width = st.number_input(
                                    "Width (mm)",
                                    value=clad_width_val,
                                    key=f"{canopy_key}_clad_width",
                                    min_value=0
                                )
                            
                            with clad_col2:
                                # Safe conversion to int
                                clad_height_val = canopy.get('wall_cladding', {}).get('height', 2100)
                                if clad_height_val is None or clad_height_val == '':
                                    clad_height_val = 2100
                                try:
                                    clad_height_val = int(clad_height_val)
                                except (ValueError, TypeError):
                                    clad_height_val = 2100
                                    
                                clad_height = st.number_input(
                                    "Height (mm)",
                                    value=clad_height_val,
                                    key=f"{canopy_key}_clad_height",
                                    min_value=0
                                )
                            
                            with clad_col3:
                                current_positions = canopy.get('wall_cladding', {}).get('position', [])
                                if isinstance(current_positions, str):
                                    current_positions = [current_positions] if current_positions else []
                                elif current_positions is None:
                                    current_positions = []
                                
                                clad_positions = st.multiselect(
                                    "Position",
                                    options=["rear", "left hand", "right hand"],
                                    default=current_positions,
                                    key=f"{canopy_key}_clad_pos"
                                )
                            
                            # Update wall cladding in session state
                            wall_cladding_data = {
                                "type": "Custom",
                                "width": clad_width,
                                "height": clad_height,
                                "position": clad_positions
                            }
                            st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['wall_cladding'] = wall_cladding_data
                        else:
                            # Update to no wall cladding
                            st.session_state.levels[level_idx]['areas'][area_idx]['canopies'][canopy_idx]['wall_cladding'] = {
                                "type": "None",
                                "width": None,
                                "height": None,
                                "position": None
                            }
                        
                        st.markdown("---")
    
    # Project Summary Section
    if st.session_state.project_info or st.session_state.levels:
        st.markdown("---")
        st.markdown("## 📊 Project Summary")
        
        # Project Information Summary
        if st.session_state.project_info:
            with st.expander("📋 Project Information", expanded=True):
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown("**General Info:**")
                    if st.session_state.project_info.get('project_name'):
                        st.write(f"Project: {st.session_state.project_info['project_name']}")
                    if st.session_state.project_info.get('project_number'):
                        st.write(f"Number: {st.session_state.project_info['project_number']}")
                    if st.session_state.project_info.get('customer'):
                        st.write(f"Customer: {st.session_state.project_info['customer']}")
                
                with col2:
                    st.markdown("**Company Details:**")
                    if st.session_state.project_info.get('company'):
                        st.write(f"Company: {st.session_state.project_info['company']}")
                    if st.session_state.project_info.get('estimator'):
                        st.write(f"Estimator: {st.session_state.project_info['estimator']}")
                    if st.session_state.project_info.get('sales_contact'):
                        st.write(f"Sales: {st.session_state.project_info['sales_contact']}")
                
                with col3:
                    st.markdown("**Location Info:**")
                    if st.session_state.project_info.get('project_location'):
                        st.write(f"Project Location: {st.session_state.project_info['project_location']}")
                    if st.session_state.project_info.get('delivery_location'):
                        st.write(f"Delivery: {st.session_state.project_info['delivery_location']}")
                    if st.session_state.project_info.get('date'):
                        st.write(f"Date: {st.session_state.project_info['date']}")
                    if st.session_state.project_info.get('contract_option') is not None:
                        st.write(f"Contract Sheets: {'Yes' if st.session_state.project_info.get('contract_option', False) else 'No'}")
        
        # Structure Summary
        if st.session_state.levels:
            with st.expander("🏢 Project Structure", expanded=True):
                total_areas = sum(len(level.get('areas', [])) for level in st.session_state.levels)
                total_canopies = sum(
                    len(area.get('canopies', [])) 
                    for level in st.session_state.levels 
                    for area in level.get('areas', [])
                )
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Levels", len(st.session_state.levels))
                with col2:
                    st.metric("Areas", total_areas)
                with col3:
                    st.metric("Canopies", total_canopies)
                
                # Detailed breakdown
                st.markdown("---")
                for level in st.session_state.levels:
                    st.markdown(f"**{level['level_name']}**")
                    for area in level.get('areas', []):
                        canopy_count = len(area.get('canopies', []))
                        options = []
                        if area['options'].get('uvc'): options.append("UV-C")
                        if area['options'].get('recoair'): options.append("RecoAir")
                        if area['options'].get('marvel'): options.append("Marvel")
                        if area['options'].get('uv_extra_over'): options.append("UV Extra")
                        if area['options'].get('vent_clg'): options.append("VENT CLG")
                        
                        options_str = f" ({', '.join(options)})" if options else ""
                        st.write(f"  • {area['name']}: {canopy_count} canopies{options_str}")
                        
                        # Show canopy details
                        for canopy in area.get('canopies', []):
                            canopy_info = []
                            if canopy.get('reference_number'):
                                canopy_info.append(f"Ref: {canopy['reference_number']}")
                            if canopy.get('model'):
                                canopy_info.append(f"Model: {canopy['model']}")
                            if canopy.get('configuration'):
                                canopy_info.append(f"Config: {canopy['configuration']}")
                            if canopy.get('options', {}).get('fire_suppression'):
                                canopy_info.append("Fire Supp")
                            if canopy.get('options', {}).get('sdu'):
                                canopy_info.append("SDU")
                            if canopy.get('wall_cladding', {}).get('type') != 'None':
                                canopy_info.append("Wall Clad")
                            
                            if canopy_info:
                                st.write(f"    - {' | '.join(canopy_info)}")

def main():
    st.set_page_config(page_title="Halton Quotation System", page_icon="🏭", layout="wide")
    st.title("Halton Quotation System")
    
    # Initialize session state
    initialize_session_state()
    
    # Sidebar navigation
    st.sidebar.title("Navigation")
    
    page = st.sidebar.selectbox(
        "Choose a page:",
        ["Single Page Setup", "Generate Word Documents", "Create Revision"]  # "Project Setup" commented out
    )
    
    # Add project summary to sidebar
    if page == "Project Setup" and st.session_state.project_info:
        st.sidebar.markdown("---")
        st.sidebar.markdown("### 📋 Current Project")
        
        # General Project Information
        with st.sidebar.expander("General Information", expanded=True):
            if st.session_state.project_info.get('project_name'):
                st.sidebar.markdown(f"**Project Name:** {st.session_state.project_info.get('project_name')}")
            if st.session_state.project_info.get('project_number'):
                st.sidebar.markdown(f"**Project Number:** {st.session_state.project_info.get('project_number')}")
            if st.session_state.project_info.get('customer'):
                st.sidebar.markdown(f"**Customer:** {st.session_state.project_info.get('customer')}")
            if st.session_state.project_info.get('company'):
                st.sidebar.markdown(f"**Company:** {st.session_state.project_info.get('company')}")
            if st.session_state.project_info.get('project_location'):
                st.sidebar.markdown(f"**Location:** {st.session_state.project_info.get('project_location')}")
            if st.session_state.project_info.get('estimator'):
                st.sidebar.markdown(f"**Estimator:** {st.session_state.project_info.get('estimator')}")
            if st.session_state.project_info.get('date'):
                st.sidebar.markdown(f"**Date:** {st.session_state.project_info.get('date')}")
            if st.session_state.project_info.get('revision'):
                st.sidebar.markdown(f"**Revision:** {st.session_state.project_info.get('revision')}")
            if st.session_state.project_info.get('delivery_location'):
                st.sidebar.markdown(f"**Delivery:** {st.session_state.project_info.get('delivery_location')}")
            if st.session_state.project_info.get('contract_option'):
                st.sidebar.markdown(f"**Contract Sheets:** {'Yes' if st.session_state.project_info.get('contract_option') else 'No'}")
        
        # Structure summary
        if st.session_state.levels:
            st.sidebar.markdown("---")
            st.sidebar.markdown("### 🏗️ Project Structure")
            
            # Overall stats
            total_areas = sum(len(level.get('areas', [])) for level in st.session_state.levels)
            total_canopies = sum(
                len(area.get('canopies', [])) 
                for level in st.session_state.levels 
                for area in level.get('areas', [])
            )
            st.sidebar.markdown(f"**Total:** {len(st.session_state.levels)} Levels, {total_areas} Areas, {total_canopies} Canopies")
            
            # Show level details with area options
            with st.sidebar.expander("Detailed Structure", expanded=False):
                for level in st.session_state.levels:
                    st.sidebar.markdown(f"**{level['level_name']}:**")
                    for area in level.get('areas', []):
                        canopy_count = len(area.get('canopies', []))
                        st.sidebar.markdown(f"**• {area['name']}** ({canopy_count} canopies)")
                        
                        # Show area options
                        options = []
                        if area.get('options', {}).get('uvc'):
                            options.append("UV-C")
                        if area.get('options', {}).get('recoair'):
                            options.append("RecoAir")
                        if area.get('options', {}).get('marvel'):
                            options.append("Marvel")
                        if area.get('options', {}).get('uv_extra_over'):
                            options.append("UV Extra Over")
                        if area.get('options', {}).get('vent_clg'):
                            options.append("VENT CLG")
                        
                        if options:
                            st.sidebar.markdown(f"  Options: {', '.join(options)}")
                        
                        # Show canopy details
                        if canopy_count > 0:
                            for canopy in area.get('canopies', []):
                                canopy_info = []
                                if canopy.get('reference_number'):
                                    canopy_info.append(f"{canopy['reference_number']}")
                                if canopy.get('model'):
                                    canopy_info.append(f"{canopy['model']}")
                                if canopy.get('options', {}).get('fire_suppression'):
                                    canopy_info.append("FS")
                                if canopy.get('options', {}).get('sdu'):
                                    sdu_text = "SDU"
                                    if canopy.get('sdu_item_number'):
                                        sdu_text += f" ({canopy['sdu_item_number']})"
                                    canopy_info.append(sdu_text)
                                
                                if canopy_info:
                                    st.sidebar.markdown(f"    - {' | '.join(canopy_info)}")
                    
                    st.sidebar.markdown("")  # Add spacing between levels
    
    # Page routing
    if False and page == "Project Setup":  # Commented out Project Setup page for now
        # Template Selection
        st.markdown("### Cost Sheet Template Selection")
        template_options = {
            "Cost Sheet R19.2 Jun 2025": "templates/excel/Cost Sheet R19.2 Jun 2025.xlsx",
            "Cost Sheet R19.1 May 2025": "templates/excel/Cost Sheet R19.1 May 2025.xlsx",
            "Cost Sheet R18.1 (Legacy)": "templates/excel/Halton Cost Sheet Jan 2025.xlsx"
        }
        
        # Initialize template selection in session state
        if "selected_template" not in st.session_state:
            st.session_state.selected_template = "Cost Sheet R19.2 Jun 2025"  # Default to 19.2
        
        # Ensure the selected template is in the available options
        template_keys = list(template_options.keys())
        if st.session_state.selected_template not in template_keys:
            # If the session template is not available, default to the first option
            st.session_state.selected_template = template_keys[0]
            st.warning(f"⚠️ Previous template version not available. Defaulted to {template_keys[0]}")
        
        selected_template = st.selectbox(
            "Choose Cost Sheet Template:",
            options=template_keys,
            index=template_keys.index(st.session_state.selected_template),
            key="template_selector",
            help="Select which version of the cost sheet template to use for this project"
        )
        
        # Update session state when selection changes
        if selected_template != st.session_state.selected_template:
            st.session_state.selected_template = selected_template
            st.rerun()
        
        # Store the template path for use in Excel operations
        st.session_state.template_path = template_options[selected_template]
        
        # Display template status
        template_path = template_options[selected_template]
        if os.path.exists(template_path) or os.path.exists(f"../{template_path}"):
            st.success(f"✅ Using template: {selected_template}")
        else:
            st.warning(f"⚠️  Template file not found: {template_path}")
            st.info("Please ensure the template file exists before generating Excel files.")
        
        st.markdown("---")
        
        # Excel Upload Feature
        st.markdown("### Quick Start from Existing Project")
        st.markdown("Upload an existing Excel file to auto-populate form fields")
        
        uploaded_file = st.file_uploader(
            "Choose an Excel file", 
            type=['xlsx'],
            help="Upload a previous project Excel file to automatically fill in the form"
        )
        
        # Process uploaded file with AI loading effect
        if uploaded_file is not None and not st.session_state.upload_success:
            try:
                # Save uploaded file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    temp_path = tmp_file.name
                
                # Simplified loading effect without rainbow border
                with st.container():
                    # Create columns for better layout
                    col1, col2 = st.columns([2, 1])
                    
                    with col1:
                        # Progress container
                        progress_container = st.container()
                        status_container = st.container()
                    
                    with col2:
                        # Animation placeholder
                        ai_placeholder = st.empty()
                    
                    # Loading animation frames
                    ai_frames = [
                        "Initializing AI...",
                        "Analyzing Excel structure...",
                        "Reading project metadata...",
                        "Extracting canopy data...",
                        "Processing fire suppression...",
                        "Analyzing lighting options...",
                        "Calculating pricing...",
                        "Finalizing extraction..."
                    ]
                    
                    # Progress bar
                    progress_bar = progress_container.progress(0)
                    status_text = status_container.empty()
                    
                    # Animate the loading process
                    for i, frame in enumerate(ai_frames):
                        progress = (i + 1) / len(ai_frames)
                        progress_bar.progress(progress)
                        status_text.markdown(f"**{frame}**")
                        ai_placeholder.markdown(f"### {frame}")
                        time.sleep(0.5)  # Pause for effect
                    
                    # Final processing
                    status_text.markdown("**Processing project data...**")
                    ai_placeholder.markdown("### Processing project data...")
                    
                    # Actually extract the data
                    extracted_data = read_excel_project_data(temp_path)
                    
                    # Success animation
                    progress_bar.progress(1.0)
                    status_text.markdown("**Extraction Complete!**")
                    ai_placeholder.markdown("### Extraction Complete!")
                    time.sleep(1)
                    
                    # Clear loading animation
                    progress_container.empty()
                    status_container.empty()
                    ai_placeholder.empty()
                    
                    # Store extracted data and immediately populate session state
                    st.session_state.uploaded_project_data = extracted_data
                    st.session_state.upload_success = True
                    
                    # Immediately populate all session state with uploaded data
                    populate_session_state_from_uploaded_data(extracted_data)
                    
                    # Success message
                    st.success("Project data extracted successfully! Form fields have been auto-populated.")
                    
                    # Show extracted project summary
                    with st.expander("Extracted Project Summary", expanded=True):
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            st.markdown(f"**Project:** {extracted_data.get('project_name', 'N/A')}")
                            st.markdown(f"**Customer:** {extracted_data.get('customer', 'N/A')}")
                            st.markdown(f"**Company:** {extracted_data.get('company', 'N/A')}")
                        
                        with col2:
                            st.markdown(f"**Project Number:** {extracted_data.get('project_number', 'N/A')}")
                            st.markdown(f"**Estimator:** {extracted_data.get('estimator', 'N/A')}")
                            st.markdown(f"**Date:** {extracted_data.get('date', 'N/A')}")
                        
                        with col3:
                            st.markdown(f"**Revision:** {extracted_data.get('revision', 'Initial') or 'Initial'}")
                            total_levels = len(extracted_data.get('levels', []))
                            total_areas = sum(len(level.get('areas', [])) for level in extracted_data.get('levels', []))
                            total_canopies = sum(len(area.get('canopies', [])) for level in extracted_data.get('levels', []) for area in level.get('areas', []))
                            st.markdown(f"**Levels:** {total_levels}")
                            st.markdown(f"**Areas:** {total_areas}")
                            st.markdown(f"**Canopies:** {total_canopies}")
                
                # Clean up temp file
                os.unlink(temp_path)
                
            except Exception as e:
                st.error(f"Error extracting data from Excel file: {str(e)}")
                st.session_state.uploaded_project_data = None
                st.session_state.upload_success = False
                # Clean up temp file if it exists
                try:
                    os.unlink(temp_path)
                except:
                    pass
        
        # Add a button to clear uploaded data
        if st.session_state.upload_success:
            if st.button("Clear Uploaded Data", help="Clear uploaded data and start fresh"):
                # Clear uploaded data flags
                st.session_state.uploaded_project_data = None
                st.session_state.upload_success = False
                st.session_state.current_step = 1  # Reset to step 1
                
                # Clear all project data
                st.session_state.project_info = {}
                st.session_state.levels = []
                
                # Clear all form state variables
                form_fields_to_clear = [
                    'project_name_state', 'customer_state', 'location_state', 
                    'project_number_state', 'revision_state', 'custom_company_name_state', 
                    'custom_company_address_state'
                ]
                for field in form_fields_to_clear:
                    if field in st.session_state:
                        del st.session_state[field]
                
                st.rerun()
        
        st.markdown("---")
        
        # Multi-step process
        # Step routing
        if st.session_state.current_step == 1:
            step1_project_information()
        elif st.session_state.current_step == 2:
            step2_project_structure()
        elif st.session_state.current_step == 3:
            step3_canopy_configuration()
        # Comment out step 4 for now
        # elif st.session_state.current_step == 4:
        #     step4_review_and_generate()
        
        st.markdown("---")
        
        # Navigation buttons
        navigation_buttons()
        
    if page == "Single Page Setup":
        single_page_project_builder()
        
    elif page == "Generate Word Documents":
        word_generation_page()
        
    elif page == "Create Revision":
        revision_page()

if __name__ == "__main__":
    main()