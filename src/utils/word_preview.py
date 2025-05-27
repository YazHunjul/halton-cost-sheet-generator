"""
Word document preview utilities for the Halton Cost Sheet Generator.
Provides functionality to preview Word documents in the web browser.
"""

import streamlit as st
from docx import Document
import tempfile
import os
import sys
import subprocess
import shutil
from typing import Optional, Dict, Any
import base64

def clear_preview_cache():
    """
    Clear any cached preview modules to ensure fresh imports.
    Useful when there are import or version conflicts.
    """
    modules_to_clear = [k for k in sys.modules.keys() if 'word_preview' in k or 'pypandoc' in k]
    for module in modules_to_clear:
        if module in sys.modules:
            del sys.modules[module]

def setup_pandoc_path():
    """
    Try to find and set up pandoc path for pypandoc.
    This helps when pandoc is installed but not in the expected PATH.
    """
    try:
        # First, try to find pandoc using which/where
        pandoc_path = shutil.which('pandoc')
        if pandoc_path:
            return pandoc_path
        
        # Try common installation paths
        common_paths = [
            '/usr/local/bin/pandoc',
            '/opt/homebrew/bin/pandoc',
            '/usr/bin/pandoc',
            '/usr/local/pandoc/bin/pandoc'
        ]
        
        for path in common_paths:
            if os.path.exists(path):
                return path
        
        return None
    except Exception:
        return None

def check_preview_requirements() -> Dict[str, bool]:
    """
    Check what preview capabilities are available.
    
    Returns:
        Dict with availability of different preview methods
    """
    capabilities = {
        'basic_preview': True,  # Always available with python-docx
        'advanced_preview': False,  # Requires pypandoc
        'table_preservation': True,  # Available with both methods
        'pandoc_version': None
    }
    
    try:
        import pypandoc
        
        # Try to set up pandoc path if needed
        pandoc_path = setup_pandoc_path()
        if pandoc_path:
            os.environ['PYPANDOC_PANDOC'] = pandoc_path
        
        # Test if pypandoc actually works by trying to get version
        try:
            version_info = pypandoc.get_pandoc_version()
            capabilities['advanced_preview'] = True
            capabilities['pandoc_version'] = str(version_info)
        except Exception as e:
            # Try alternative methods to detect pandoc
            try:
                # Try to run a simple conversion test
                import tempfile
                with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as tmp:
                    tmp.write("# Test")
                    tmp_path = tmp.name
                
                # Try a simple conversion
                pypandoc.convert_file(tmp_path, 'html')
                
                # If we get here, pypandoc works but version detection failed
                capabilities['advanced_preview'] = True
                capabilities['pandoc_version'] = "Available (version detection failed)"
                
                # Clean up
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
                    
            except Exception as e2:
                # pypandoc is installed but pandoc binary is not available
                capabilities['advanced_preview'] = False
                capabilities['pandoc_version'] = "Pandoc binary not found"
            
    except ImportError:
        capabilities['pandoc_version'] = "pypandoc not installed"
    
    return capabilities

def extract_text_content(docx_path: str) -> str:
    """
    Extract plain text content from a Word document.
    
    Args:
        docx_path (str): Path to the Word document
        
    Returns:
        str: Plain text content
    """
    try:
        doc = Document(docx_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        return '\n\n'.join(text_content)
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def convert_docx_to_html_simple(docx_path: str) -> str:
    """
    Convert DOCX to HTML using python-docx with enhanced table support.
    
    Args:
        docx_path (str): Path to the Word document
        
    Returns:
        str: HTML content with preserved tables
    """
    try:
        doc = Document(docx_path)
        html_content = []
        
        # Add CSS for professional styling
        html_content.append("""
        <style>
            body { font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }
            h1, h2, h3 { color: #2c3e50; margin-top: 30px; margin-bottom: 15px; }
            h1 { font-size: 24px; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
            h2 { font-size: 20px; color: #34495e; }
            h3 { font-size: 16px; color: #7f8c8d; }
            p { margin-bottom: 12px; }
            table { 
                border-collapse: collapse; 
                width: 100%; 
                margin: 20px 0; 
                box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            }
            th, td { 
                border: 1px solid #bdc3c7; 
                padding: 12px 8px; 
                text-align: left; 
                vertical-align: top;
            }
            th { 
                background: linear-gradient(135deg, #74b9ff 0%, #0984e3 100%);
                font-weight: bold; 
                color: white;
            }
            tr:nth-child(even) { background-color: #f8f9fa; }
            tr:hover { background-color: #e8f4f8; }
            .table-container { overflow-x: auto; margin: 20px 0; }
        </style>
        """)
        
        # Process all document elements in order
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                # Find the corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:
                            # Detect heading styles
                            if para.style.name.startswith('Heading'):
                                level = para.style.name.replace('Heading ', '')
                                try:
                                    level_num = int(level)
                                    html_content.append(f'<h{min(level_num, 6)}>{text}</h{min(level_num, 6)}>')
                                except:
                                    html_content.append(f'<h3>{text}</h3>')
                            else:
                                html_content.append(f'<p>{text}</p>')
                        break
            
            elif element.tag.endswith('tbl'):  # Table
                # Find the corresponding table object
                for table in doc.tables:
                    if table._element == element:
                        html_content.append('<div class="table-container">')
                        html_content.append('<table>')
                        
                        for i, row in enumerate(table.rows):
                            html_content.append('<tr>')
                            for cell in row.cells:
                                cell_text = cell.text.strip() or '&nbsp;'
                                # First row is typically headers
                                tag = 'th' if i == 0 else 'td'
                                html_content.append(f'<{tag}>{cell_text}</{tag}>')
                            html_content.append('</tr>')
                        
                        html_content.append('</table>')
                        html_content.append('</div>')
                        break
        
        return '\n'.join(html_content)
        
    except Exception as e:
        return f"<p>Error converting document: {str(e)}</p>"

def convert_docx_to_html_advanced(docx_path: str) -> str:
    """
    Convert DOCX to HTML using pypandoc with enhanced table styling.
    
    Args:
        docx_path (str): Path to the Word document
        
    Returns:
        str: HTML content with enhanced styling
    """
    try:
        import pypandoc
        
        # Try advanced conversion with multiple fallback strategies
        html_content = None
        error_messages = []
        
        # Strategy 1: Try with --standalone option
        try:
            html_content = pypandoc.convert_file(
                docx_path, 
                'html',
                extra_args=['--standalone']
            )
        except Exception as e1:
            error_messages.append(f"Strategy 1 (--standalone): {str(e1)}")
            
            # Strategy 2: Try without any extra args
            try:
                html_content = pypandoc.convert_file(docx_path, 'html')
            except Exception as e2:
                error_messages.append(f"Strategy 2 (no args): {str(e2)}")
                
                # Strategy 3: Fall back to basic conversion
                try:
                    return convert_docx_to_html_simple(docx_path)
                except Exception as e3:
                    error_messages.append(f"Strategy 3 (basic fallback): {str(e3)}")
                    return f"<p>All conversion strategies failed:<br>{'<br>'.join(error_messages)}</p>"
        
        # If we got here, one of the strategies worked
        if html_content is None:
            return convert_docx_to_html_simple(docx_path)
        
        # Enhanced CSS for better table styling
        enhanced_css = """
        <style>
            body { 
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
                line-height: 1.6; 
                margin: 20px; 
                color: #333;
            }
            h1, h2, h3, h4, h5, h6 { 
                color: #2c3e50; 
                margin-top: 30px; 
                margin-bottom: 15px; 
            }
            h1 { 
                font-size: 28px; 
                border-bottom: 3px solid #3498db; 
                padding-bottom: 10px; 
            }
            h2 { 
                font-size: 22px; 
                color: #34495e; 
                border-bottom: 1px solid #bdc3c7;
                padding-bottom: 5px;
            }
            h3 { font-size: 18px; color: #7f8c8d; }
            p { margin-bottom: 15px; }
            table { 
                border-collapse: collapse; 
                width: 100%; 
                margin: 25px 0; 
                box-shadow: 0 4px 8px rgba(0,0,0,0.1);
                border-radius: 8px;
                overflow: hidden;
            }
            th, td { 
                border: 1px solid #ddd; 
                padding: 15px 12px; 
                text-align: left; 
                vertical-align: top;
            }
            th { 
                background: linear-gradient(135deg, #74b9ff 0%, #0984e3 100%);
                color: white;
                font-weight: 600;
                text-transform: uppercase;
                font-size: 12px;
                letter-spacing: 0.5px;
            }
            tr:nth-child(even) { background-color: #f8f9fa; }
            tr:hover { 
                background-color: #e3f2fd; 
                transition: background-color 0.3s ease;
            }
            .table-container { 
                overflow-x: auto; 
                margin: 25px 0; 
                border-radius: 8px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }
            strong { color: #2c3e50; }
            em { color: #7f8c8d; }
        </style>
        """
        
        # Inject enhanced CSS
        if '<head>' in html_content:
            html_content = html_content.replace('<head>', f'<head>{enhanced_css}')
        else:
            html_content = f'{enhanced_css}\n{html_content}'
        
        return html_content
        
    except ImportError:
        # Fallback to basic conversion if pypandoc not available
        return convert_docx_to_html_simple(docx_path)
    except Exception as e:
        # Final fallback with error information
        return f"<p>Error with advanced conversion: {str(e)}<br>Falling back to basic conversion...</p>" + convert_docx_to_html_simple(docx_path)

def preview_word_document(docx_path: str, use_advanced: bool = True) -> str:
    """
    Generate HTML preview of a Word document with table preservation.
    
    Args:
        docx_path (str): Path to the Word document
        use_advanced (bool): Whether to use advanced conversion (pypandoc)
        
    Returns:
        str: HTML content for preview
    """
    capabilities = check_preview_requirements()
    
    if use_advanced and capabilities['advanced_preview']:
        html_content = convert_docx_to_html_advanced(docx_path)
    else:
        html_content = convert_docx_to_html_simple(docx_path)
    
    # Wrap in a container with scrolling
    preview_html = f"""
    <div style="
        max-height: 600px; 
        overflow-y: auto; 
        border: 1px solid #ddd; 
        border-radius: 8px; 
        padding: 20px; 
        background-color: white;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
    ">
        {html_content}
    </div>
    """
    
    return preview_html

def preview_with_download(docx_path: str, filename: str) -> None:
    """
    Display document preview with download option in Streamlit.
    
    Args:
        docx_path (str): Path to the Word document
        filename (str): Filename for download
    """
    st.subheader("📄 Document Preview")
    
    # Check capabilities
    capabilities = check_preview_requirements()
    
    # Preview options
    col1, col2 = st.columns([3, 1])
    
    with col2:
        use_advanced = st.checkbox(
            "Enhanced Preview", 
            value=capabilities['advanced_preview'],
            disabled=not capabilities['advanced_preview'],
            help="Uses pypandoc for better formatting (if available)"
        )
        
        if not capabilities['advanced_preview']:
            st.info("💡 Install pypandoc for enhanced preview")
        elif capabilities['pandoc_version']:
            st.caption(f"Pandoc version: {capabilities['pandoc_version']}")
    
    with col1:
        st.write("**Preview Mode:**", "Enhanced" if use_advanced else "Basic")
        if capabilities['table_preservation']:
            st.write("✅ Table preservation enabled")
    
    # Generate and display preview
    try:
        with st.spinner("Generating preview..."):
            preview_html = preview_word_document(docx_path, use_advanced)
        
        # Display preview
        st.components.v1.html(preview_html, height=650, scrolling=True)
        
        # Download button
        st.markdown("---")
        with open(docx_path, "rb") as file:
            st.download_button(
                label="📥 Download Document",
                data=file.read(),
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
        
        # Preview stats
        try:
            doc = Document(docx_path)
            table_count = len(doc.tables)
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📄 Paragraphs", paragraph_count)
            with col2:
                st.metric("📊 Tables", table_count)
            with col3:
                file_size = os.path.getsize(docx_path)
                st.metric("💾 File Size", f"{file_size/1024:.1f} KB")
                
        except Exception as e:
            st.write(f"Preview stats unavailable: {str(e)}")
            
    except Exception as e:
        st.error(f"❌ Error generating preview: {str(e)}")
        
        # Fallback: still offer download
        st.markdown("---")
        with open(docx_path, "rb") as file:
            st.download_button(
                label="📥 Download Document (Preview Failed)",
                data=file.read(),
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ) 